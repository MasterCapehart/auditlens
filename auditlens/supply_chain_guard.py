"""
AuditLens Supply Chain Security Suite

Provides comprehensive supply chain security auditing including:
- SBOM diff analysis (CycloneDX, SPDX)
- Typosquatting detection
- Dependency confusion scanning
- License compliance checking
- Supply chain attack detection
- SBOM integrity verification
- Multi-format reporting (HTML, JSON, DOCX, XLSX)

Usage:
    from auditlens.supply_chain_guard import audit_supply_chain

    report = audit_supply_chain('./project', severity='MEDIUM')
    generate_supply_chain_report(report, format='html', output_path='sc_report.html')
"""

from __future__ import annotations

import json
import os
import re
import sqlite3
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field, asdict
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple
from urllib.parse import quote

try:
    import requests
    _REQUESTS_AVAILABLE = True
except ImportError:
    _REQUESTS_AVAILABLE = False


# ── Data Structures ────────────────────────────────────────────────────────────

@dataclass
class Component:
    """Represents a software component in the supply chain."""
    name: str
    version: str
    ecosystem: str  # pypi, npm, maven, etc.
    purl: str = ""
    license: str = ""
    risk_score: float = 0.0
    metadata: Dict = field(default_factory=dict)

    def __post_init__(self):
        if not self.purl:
            self.purl = self._generate_purl()

    def _generate_purl(self) -> str:
        """Generate Package URL (purl) from component info."""
        eco_map = {
            'pypi': 'pkg:pypi/',
            'npm': 'pkg:npm/',
            'maven': 'pkg:maven/',
            'gem': 'pkg:gem/',
            'cargo': 'pkg:cargo/',
        }
        prefix = eco_map.get(self.ecosystem.lower(), f'pkg:{self.ecosystem}/')
        return f"{prefix}{quote(self.name)}@{quote(self.version)}"


@dataclass
class ComponentChange:
    """Represents a change in component version."""
    name: str
    old_version: str
    new_version: str
    change_type: str  # upgrade, downgrade, major_change
    risk: str = "LOW"  # LOW, MEDIUM, HIGH, CRITICAL
    reason: str = ""


@dataclass
class DiffResult:
    """Result of comparing two SBOMs."""
    added: List[Component] = field(default_factory=list)
    removed: List[Component] = field(default_factory=list)
    modified: List[ComponentChange] = field(default_factory=list)
    unchanged: int = 0
    risk_score: float = 0.0
    timestamp: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())


@dataclass
class TyposquattingMatch:
    """Detected typosquatting candidate."""
    suspicious_package: str
    legitimate_package: str
    similarity_score: float
    detection_method: str  # levenshtein, confusables, keyboard
    ecosystem: str


@dataclass
class IntegrityReport:
    """SBOM integrity verification result."""
    undeclared_deps: List[Dict] = field(default_factory=list)
    phantom_components: List[Dict] = field(default_factory=list)
    checksum_mismatches: List[Dict] = field(default_factory=list)
    integrity_score: float = 100.0


@dataclass
class SupplyChainReport:
    """Complete supply chain audit report."""
    sbom_diff: Optional[DiffResult] = None
    license_findings: List[Dict] = field(default_factory=list)
    confusion_findings: List[Dict] = field(default_factory=list)
    typosquatting_findings: List[Dict] = field(default_factory=list)
    integrity_issues: List[Dict] = field(default_factory=list)
    risk_summary: Dict = field(default_factory=dict)
    score: float = 0.0
    recommendations: List[str] = field(default_factory=list)
    generated_at: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())


# ── SBOM Differ ────────────────────────────────────────────────────────────────

class SBOMDiffer:
    """Compares two SBOMs and detects suspicious changes."""

    def __init__(self):
        self._parsers = {
            'cyclonedx': self.parse_cyclonedx,
            'spdx': self.parse_spdx,
        }

    def diff(self, baseline_path: str, current_path: str, format: str = 'cyclonedx') -> DiffResult:
        """Compare two SBOMs and return differences."""
        if format not in self._parsers:
            raise ValueError(f"Unsupported SBOM format: {format}. Use 'cyclonedx' or 'spdx'.")

        parser = self._parsers[format]
        baseline_components = parser(baseline_path) if os.path.exists(baseline_path) else {}
        current_components = parser(current_path)

        baseline_purls = set(baseline_components.keys())
        current_purls = set(current_components.keys())

        added = [current_components[p] for p in (current_purls - baseline_purls)]
        removed = [baseline_components[p] for p in (baseline_purls - current_purls)]

        modified = []
        for purl in (baseline_purls & current_purls):
            old_comp = baseline_components[purl]
            new_comp = current_components[purl]
            if old_comp.version != new_comp.version:
                change = self._analyze_version_change(old_comp, new_comp)
                modified.append(change)

        result = DiffResult(
            added=added,
            removed=removed,
            modified=modified,
            unchanged=len(baseline_purls & current_purls) - len(modified)
        )

        result.risk_score = self._calculate_diff_risk(result)
        return result

    def parse_cyclonedx(self, path: str) -> Dict[str, Component]:
        """Parse CycloneDX SBOM and return components indexed by PURL."""
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        components = {}
        for comp in data.get('components', []):
            component = self._normalize_cyclonedx_component(comp)
            if component:
                components[component.purl] = component

        return components

    def parse_spdx(self, path: str) -> Dict[str, Component]:
        """Parse SPDX SBOM and return components indexed by PURL."""
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        components = {}
        for pkg in data.get('packages', []):
            component = self._normalize_spdx_component(pkg)
            if component:
                components[component.purl] = component

        return components

    def _normalize_cyclonedx_component(self, raw: dict) -> Optional[Component]:
        """Convert CycloneDX component to normalized Component."""
        name = raw.get('name', '').strip()
        version = raw.get('version', '').strip()
        if not name or not version:
            return None

        purl = raw.get('purl', '')
        ecosystem = self._extract_ecosystem_from_purl(purl) or raw.get('type', 'library')

        licenses = []
        for lic in raw.get('licenses', []):
            if 'license' in lic:
                licenses.append(lic['license'].get('id') or lic['license'].get('name', ''))

        return Component(
            name=name,
            version=version,
            ecosystem=ecosystem,
            purl=purl,
            license=', '.join(licenses) if licenses else 'Unknown',
            metadata=raw
        )

    def _normalize_spdx_component(self, raw: dict) -> Optional[Component]:
        """Convert SPDX package to normalized Component."""
        name = raw.get('name', '').strip()
        version = raw.get('versionInfo', '').strip()
        if not name or not version:
            return None

        # Extract PURL from externalRefs
        purl = ''
        for ref in raw.get('externalRefs', []):
            if ref.get('referenceType') == 'purl':
                purl = ref.get('referenceLocator', '')
                break

        ecosystem = self._extract_ecosystem_from_purl(purl) or 'unknown'

        license = raw.get('licenseConcluded', raw.get('licenseDeclared', 'NOASSERTION'))

        return Component(
            name=name,
            version=version,
            ecosystem=ecosystem,
            purl=purl,
            license=license,
            metadata=raw
        )

    def _extract_ecosystem_from_purl(self, purl: str) -> str:
        """Extract ecosystem type from Package URL."""
        if not purl:
            return ''
        match = re.match(r'pkg:([^/]+)/', purl)
        return match.group(1) if match else ''

    def _analyze_version_change(self, old: Component, new: Component) -> ComponentChange:
        """Analyze version change and determine risk."""
        old_parts = self._parse_semver(old.version)
        new_parts = self._parse_semver(new.version)

        if old_parts and new_parts:
            old_maj, old_min, old_patch = old_parts
            new_maj, new_min, new_patch = new_parts

            if new_maj < old_maj:
                return ComponentChange(
                    name=old.name,
                    old_version=old.version,
                    new_version=new.version,
                    change_type='downgrade',
                    risk='HIGH',
                    reason='Major version downgrade detected — potential supply chain attack indicator'
                )
            elif new_maj == old_maj and new_min < old_min:
                return ComponentChange(
                    name=old.name,
                    old_version=old.version,
                    new_version=new.version,
                    change_type='downgrade',
                    risk='MEDIUM',
                    reason='Minor version downgrade — unusual for production dependencies'
                )
            elif new_maj > old_maj:
                return ComponentChange(
                    name=old.name,
                    old_version=old.version,
                    new_version=new.version,
                    change_type='major_change',
                    risk='MEDIUM',
                    reason='Major version upgrade — review breaking changes and changelogs'
                )

        return ComponentChange(
            name=old.name,
            old_version=old.version,
            new_version=new.version,
            change_type='upgrade',
            risk='LOW',
            reason='Version updated'
        )

    def _parse_semver(self, version: str) -> Optional[Tuple[int, int, int]]:
        """Parse semantic version string into (major, minor, patch)."""
        match = re.match(r'^v?(\d+)\.(\d+)\.(\d+)', version)
        if match:
            return int(match.group(1)), int(match.group(2)), int(match.group(3))
        return None

    def _calculate_diff_risk(self, result: DiffResult) -> float:
        """Calculate overall risk score for SBOM diff (0-100)."""
        score = 0.0

        # New components are moderate risk
        score += len(result.added) * 5

        # Removed components are low risk
        score += len(result.removed) * 2

        # Modified components weighted by their risk
        risk_weights = {'CRITICAL': 20, 'HIGH': 15, 'MEDIUM': 10, 'LOW': 5}
        for change in result.modified:
            score += risk_weights.get(change.risk, 5)

        return min(score, 100.0)


# ── Typosquatting Detector ─────────────────────────────────────────────────────

class TyposquattingDetector:
    """Detects potentially malicious typosquatting packages."""

    def __init__(self, threshold: float = 0.85):
        self.threshold = threshold
        self._popular_packages = {}
        self._cache_dir = Path.home() / '.auditlens' / 'supply_chain_cache'
        self._cache_dir.mkdir(parents=True, exist_ok=True)

    def scan_project(self, project_path: str) -> List[TyposquattingMatch]:
        """Scan project for potential typosquatting packages."""
        findings = []

        # Collect all dependencies
        deps = self._collect_dependencies(project_path)

        for dep_name, ecosystem in deps:
            self._load_popular_packages(ecosystem)
            matches = self._check_typosquatting(dep_name, ecosystem)
            findings.extend(matches)

        return findings

    def _collect_dependencies(self, project_path: str) -> List[Tuple[str, str]]:
        """Collect all dependencies from project manifests."""
        deps = []

        # Python
        req_files = Path(project_path).rglob('requirements*.txt')
        for req_file in req_files:
            deps.extend(self._parse_requirements_txt(req_file, 'pypi'))

        # Node.js
        package_json_files = Path(project_path).rglob('package.json')
        for pkg_file in package_json_files:
            deps.extend(self._parse_package_json(pkg_file, 'npm'))

        return deps

    def _parse_requirements_txt(self, file_path: Path, ecosystem: str) -> List[Tuple[str, str]]:
        """Extract package names from requirements.txt."""
        deps = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith(('#', '-')):
                        continue
                    match = re.match(r'^([A-Za-z0-9][A-Za-z0-9._-]*)', line)
                    if match:
                        deps.append((match.group(1), ecosystem))
        except Exception:
            pass
        return deps

    def _parse_package_json(self, file_path: Path, ecosystem: str) -> List[Tuple[str, str]]:
        """Extract package names from package.json."""
        deps = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            for section in ('dependencies', 'devDependencies'):
                for pkg_name in data.get(section, {}).keys():
                    deps.append((pkg_name, ecosystem))
        except Exception:
            pass
        return deps

    def _load_popular_packages(self, ecosystem: str):
        """Load popular package names for the given ecosystem."""
        if ecosystem in self._popular_packages:
            return

        cache_file = self._cache_dir / f'popular_{ecosystem}.json'

        # Try cache first
        if cache_file.exists():
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    self._popular_packages[ecosystem] = json.load(f)
                return
            except Exception:
                pass

        # Fallback to embedded popular packages
        self._popular_packages[ecosystem] = self._get_embedded_popular_packages(ecosystem)

    def _get_embedded_popular_packages(self, ecosystem: str) -> List[str]:
        """Return embedded list of popular packages."""
        # Top packages by ecosystem (truncated for size)
        popular = {
            'pypi': [
                'requests', 'urllib3', 'certifi', 'setuptools', 'pip', 'wheel',
                'numpy', 'pandas', 'boto3', 'django', 'flask', 'pytest',
                'pyyaml', 'click', 'jinja2', 'cryptography', 'sqlalchemy',
                'pillow', 'six', 'python-dateutil', 'pyjwt', 'redis', 'celery',
                'gunicorn', 'psycopg2', 'aiohttp', 'fastapi', 'pydantic',
            ],
            'npm': [
                'lodash', 'react', 'react-dom', 'express', 'axios', 'typescript',
                'webpack', 'eslint', 'prettier', 'jest', 'moment', 'commander',
                'chalk', 'debug', 'request', 'async', 'underscore', 'vue',
                'angular', 'jquery', 'redux', 'babel-core', 'next', 'tailwindcss',
                'socket.io', 'mongoose', 'dotenv', 'passport', 'nodemon',
            ],
        }
        return popular.get(ecosystem, [])

    def _check_typosquatting(self, package_name: str, ecosystem: str) -> List[TyposquattingMatch]:
        """Check if package name is similar to popular packages."""
        matches = []
        popular = self._popular_packages.get(ecosystem, [])

        for legit_pkg in popular:
            # Skip if exact match (it's the legitimate package)
            if package_name.lower() == legit_pkg.lower():
                continue

            # Skip if length difference too large
            if abs(len(package_name) - len(legit_pkg)) > 3:
                continue

            # Levenshtein distance
            similarity = self._compute_similarity(package_name, legit_pkg)
            if similarity >= self.threshold:
                matches.append(TyposquattingMatch(
                    suspicious_package=package_name,
                    legitimate_package=legit_pkg,
                    similarity_score=similarity,
                    detection_method='levenshtein',
                    ecosystem=ecosystem
                ))

            # Confusables (character substitution)
            if self._has_confusables(package_name, legit_pkg):
                matches.append(TyposquattingMatch(
                    suspicious_package=package_name,
                    legitimate_package=legit_pkg,
                    similarity_score=1.0,
                    detection_method='confusables',
                    ecosystem=ecosystem
                ))

        return matches

    def _compute_similarity(self, a: str, b: str) -> float:
        """Compute similarity ratio using Levenshtein distance."""
        try:
            import Levenshtein
            return Levenshtein.ratio(a.lower(), b.lower())
        except ImportError:
            # Fallback to pure Python implementation
            return self._levenshtein_ratio_pure(a.lower(), b.lower())

    def _levenshtein_ratio_pure(self, a: str, b: str) -> float:
        """Pure Python Levenshtein distance calculation."""
        if len(a) < len(b):
            a, b = b, a

        if len(b) == 0:
            return 0.0

        prev_row = list(range(len(b) + 1))
        for i, ca in enumerate(a, 1):
            curr_row = [i]
            for j, cb in enumerate(b, 1):
                insertions = prev_row[j] + 1
                deletions = curr_row[j - 1] + 1
                substitutions = prev_row[j - 1] + (ca != cb)
                curr_row.append(min(insertions, deletions, substitutions))
            prev_row = curr_row

        distance = prev_row[-1]
        max_len = max(len(a), len(b))
        return 1 - (distance / max_len)

    def _has_confusables(self, package_name: str, legit_pkg: str) -> bool:
        """Check for Unicode confusables and common character substitutions."""
        confusable_pairs = [
            ('0', 'o'), ('0', 'O'), ('1', 'l'), ('1', 'I'),
            ('_', '-'), ('rn', 'm'), ('vv', 'w'),
        ]

        for src, dst in confusable_pairs:
            if package_name.replace(src, dst).lower() == legit_pkg.lower():
                return True
            if package_name.replace(dst, src).lower() == legit_pkg.lower():
                return True

        return False


# ── Supply Chain Risk Scorer ───────────────────────────────────────────────────

class SupplyChainRiskScorer:
    """Calculates risk scores for components based on metadata."""

    def __init__(self):
        self._cache_db = Path.home() / '.auditlens' / 'supply_chain_cache.db'
        self._cache_db.parent.mkdir(parents=True, exist_ok=True)
        self._init_cache_db()

    def _init_cache_db(self):
        """Initialize SQLite cache database."""
        conn = sqlite3.connect(str(self._cache_db))
        conn.execute("""
            CREATE TABLE IF NOT EXISTS package_metadata (
                purl TEXT PRIMARY KEY,
                metadata TEXT,
                cached_at INTEGER
            )
        """)
        conn.execute("CREATE INDEX IF NOT EXISTS idx_cached_at ON package_metadata(cached_at)")
        conn.commit()
        conn.close()

    def score_component(self, name: str, version: str, ecosystem: str) -> float:
        """Calculate risk score for a component (0-100)."""
        metadata = self._get_package_metadata(name, ecosystem)
        if not metadata:
            return 50.0  # Default medium risk for unknown packages

        return self._calculate_risk_score(metadata)

    def _get_package_metadata(self, package: str, ecosystem: str) -> Optional[Dict]:
        """Fetch package metadata from registry or cache."""
        purl = f"pkg:{ecosystem}/{package}"

        # Check cache first
        conn = sqlite3.connect(str(self._cache_db))
        row = conn.execute(
            "SELECT metadata, cached_at FROM package_metadata WHERE purl = ?",
            (purl,)
        ).fetchone()

        if row:
            metadata_json, cached_at = row
            # Cache valid for 7 days
            if time.time() - cached_at < 7 * 86400:
                conn.close()
                return json.loads(metadata_json)

        conn.close()

        # Fetch from registry
        if not _REQUESTS_AVAILABLE:
            return None

        if ecosystem == 'pypi':
            metadata = self._get_pypi_metadata(package)
        elif ecosystem == 'npm':
            metadata = self._get_npm_metadata(package)
        else:
            return None

        if metadata:
            # Update cache
            conn = sqlite3.connect(str(self._cache_db))
            conn.execute(
                "INSERT OR REPLACE INTO package_metadata (purl, metadata, cached_at) VALUES (?, ?, ?)",
                (purl, json.dumps(metadata), int(time.time()))
            )
            conn.commit()
            conn.close()

        return metadata

    def _get_pypi_metadata(self, package: str) -> Optional[Dict]:
        """Fetch metadata from PyPI."""
        try:
            resp = requests.get(
                f"https://pypi.org/pypi/{package}/json",
                timeout=5
            )
            if resp.status_code == 200:
                return resp.json()
        except Exception:
            pass
        return None

    def _get_npm_metadata(self, package: str) -> Optional[Dict]:
        """Fetch metadata from npm registry."""
        try:
            resp = requests.get(
                f"https://registry.npmjs.org/{package}",
                timeout=5
            )
            if resp.status_code == 200:
                return resp.json()
        except Exception:
            pass
        return None

    def _calculate_risk_score(self, metadata: Dict) -> float:
        """Calculate risk score from package metadata."""
        score = 0.0

        # Age factor (20%): newer packages are riskier
        age_score = self._score_age(metadata)
        score += age_score * 0.2

        # Maintainer count (15%): fewer maintainers = higher risk
        maintainer_score = self._score_maintainers(metadata)
        score += maintainer_score * 0.15

        # Download popularity (25%): low downloads = higher risk
        downloads_score = self._score_downloads(metadata)
        score += downloads_score * 0.25

        # Last update (20%): stale packages are riskier
        update_score = self._score_last_update(metadata)
        score += update_score * 0.2

        # Known vulnerabilities (20%): handled separately by SCA
        score += 0 * 0.2

        return min(score, 100.0)

    def _score_age(self, metadata: Dict) -> float:
        """Score based on package age (0-100)."""
        # PyPI
        if 'info' in metadata:
            # Package creation not directly available, use first release
            releases = metadata.get('releases', {})
            if releases:
                first_version = min(releases.keys())
                first_release = releases[first_version]
                if first_release:
                    upload_time = first_release[0].get('upload_time', '')
                    if upload_time:
                        age_days = (datetime.now() - datetime.fromisoformat(upload_time.replace('Z', '+00:00'))).days
                        if age_days < 30:
                            return 80.0  # Very new, high risk
                        elif age_days < 365:
                            return 50.0  # Less than a year
                        else:
                            return 10.0  # Established package

        # npm
        if 'time' in metadata:
            created = metadata['time'].get('created', '')
            if created:
                age_days = (datetime.now(timezone.utc) - datetime.fromisoformat(created.replace('Z', '+00:00'))).days
                if age_days < 30:
                    return 80.0
                elif age_days < 365:
                    return 50.0
                else:
                    return 10.0

        return 50.0  # Unknown age

    def _score_maintainers(self, metadata: Dict) -> float:
        """Score based on number of maintainers (0-100)."""
        # npm
        if 'maintainers' in metadata:
            count = len(metadata['maintainers'])
            if count == 0:
                return 90.0
            elif count == 1:
                return 60.0
            elif count <= 3:
                return 30.0
            else:
                return 10.0

        # PyPI
        if 'info' in metadata and 'maintainer' in metadata['info']:
            maintainer = metadata['info']['maintainer']
            if not maintainer:
                return 90.0
            return 40.0

        return 50.0  # Unknown

    def _score_downloads(self, metadata: Dict) -> float:
        """Score based on download count (0-100)."""
        # npm
        if 'downloads' in metadata:
            downloads = metadata.get('downloads', 0)
            if downloads < 100:
                return 90.0
            elif downloads < 10000:
                return 60.0
            elif downloads < 100000:
                return 30.0
            else:
                return 5.0

        # PyPI - would need separate API call to pypistats
        # For now, use presence in popular packages as proxy
        return 40.0

    def _score_last_update(self, metadata: Dict) -> float:
        """Score based on last update time (0-100)."""
        # npm
        if 'time' in metadata and 'modified' in metadata['time']:
            modified = metadata['time']['modified']
            days_since = (datetime.now(timezone.utc) - datetime.fromisoformat(modified.replace('Z', '+00:00'))).days
            if days_since > 730:  # 2 years
                return 80.0
            elif days_since > 365:  # 1 year
                return 50.0
            else:
                return 10.0

        # PyPI
        if 'releases' in metadata:
            releases = metadata['releases']
            if releases:
                latest_version = max(releases.keys())
                latest_release = releases[latest_version]
                if latest_release:
                    upload_time = latest_release[0].get('upload_time', '')
                    if upload_time:
                        days_since = (datetime.now() - datetime.fromisoformat(upload_time.replace('Z', '+00:00'))).days
                        if days_since > 730:
                            return 80.0
                        elif days_since > 365:
                            return 50.0
                        else:
                            return 10.0

        return 50.0


# ── SBOM Integrity Checker ─────────────────────────────────────────────────────

class SBOMIntegrityChecker:
    """Verifies SBOM integrity and detects drift."""

    def check_integrity(self, sbom_path: str, project_path: str) -> IntegrityReport:
        """Check SBOM integrity against actual project dependencies."""
        report = IntegrityReport()

        # Parse SBOM
        differ = SBOMDiffer()
        sbom_format = self._detect_sbom_format(sbom_path)
        sbom_components = differ._parsers[sbom_format](sbom_path)

        # Parse actual project manifests
        actual_deps = self._parse_manifest_files(project_path)

        # Detect undeclared dependencies
        sbom_names = {c.name.lower() for c in sbom_components.values()}
        for dep_name, source_file in actual_deps:
            if dep_name.lower() not in sbom_names:
                report.undeclared_deps.append({
                    'name': dep_name,
                    'source': source_file,
                    'severity': 'MEDIUM',
                    'message': f'Dependency {dep_name} found in {source_file} but not in SBOM'
                })

        # Detect phantom components
        actual_names = {name.lower() for name, _ in actual_deps}
        for comp in sbom_components.values():
            if comp.name.lower() not in actual_names:
                report.phantom_components.append({
                    'name': comp.name,
                    'version': comp.version,
                    'severity': 'LOW',
                    'message': f'Component {comp.name} in SBOM but not found in project manifests'
                })

        # Calculate integrity score
        total_issues = len(report.undeclared_deps) + len(report.phantom_components)
        if total_issues == 0:
            report.integrity_score = 100.0
        else:
            # Penalize 10 points per issue, minimum 0
            report.integrity_score = max(0.0, 100.0 - (total_issues * 10))

        return report

    def _detect_sbom_format(self, sbom_path: str) -> str:
        """Detect SBOM format from file content."""
        with open(sbom_path, 'r', encoding='utf-8') as f:
            content = f.read(500)

        if 'bomFormat' in content or 'cyclonedx' in content.lower():
            return 'cyclonedx'
        elif 'spdxVersion' in content or 'SPDX' in content:
            return 'spdx'
        else:
            return 'cyclonedx'  # Default

    def _parse_manifest_files(self, project_path: str) -> List[Tuple[str, str]]:
        """Parse all manifest files and return (package_name, source_file) tuples."""
        deps = []

        # Python requirements.txt
        for req_file in Path(project_path).rglob('requirements*.txt'):
            deps.extend(self._parse_requirements(req_file))

        # Python Pipfile
        for pipfile in Path(project_path).rglob('Pipfile'):
            deps.extend(self._parse_pipfile(pipfile))

        # Node.js package.json
        for pkg_json in Path(project_path).rglob('package.json'):
            if 'node_modules' not in str(pkg_json):
                deps.extend(self._parse_package_json_deps(pkg_json))

        return deps

    def _parse_requirements(self, file_path: Path) -> List[Tuple[str, str]]:
        """Parse requirements.txt file."""
        deps = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith(('#', '-')):
                        continue
                    match = re.match(r'^([A-Za-z0-9][A-Za-z0-9._-]*)', line)
                    if match:
                        deps.append((match.group(1), str(file_path)))
        except Exception:
            pass
        return deps

    def _parse_pipfile(self, file_path: Path) -> List[Tuple[str, str]]:
        """Parse Pipfile."""
        deps = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            # Simple regex-based parsing (TOML libraries may not be available)
            for match in re.finditer(r'^([A-Za-z0-9][A-Za-z0-9._-]*)\s*=', content, re.MULTILINE):
                deps.append((match.group(1), str(file_path)))
        except Exception:
            pass
        return deps

    def _parse_package_json_deps(self, file_path: Path) -> List[Tuple[str, str]]:
        """Parse package.json dependencies."""
        deps = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            for section in ('dependencies', 'devDependencies', 'peerDependencies'):
                for pkg_name in data.get(section, {}).keys():
                    deps.append((pkg_name, str(file_path)))
        except Exception:
            pass
        return deps


# ── Supply Chain Guard (Main Orchestrator) ─────────────────────────────────────

class SupplyChainGuard:
    """Main orchestrator for supply chain security auditing."""

    def __init__(self):
        self.differ = SBOMDiffer()
        self.typosquat_detector = TyposquattingDetector()
        self.risk_scorer = SupplyChainRiskScorer()
        self.integrity_checker = SBOMIntegrityChecker()

    def audit_project(
        self,
        project_path: str,
        baseline_sbom: Optional[str] = None,
        severity: str = 'LOW'
    ) -> SupplyChainReport:
        """Execute complete supply chain audit."""
        report = SupplyChainReport()

        # SBOM diff analysis
        if baseline_sbom and os.path.exists(baseline_sbom):
            current_sbom = self._find_or_generate_sbom(project_path)
            if current_sbom:
                report.sbom_diff = self.differ.diff(baseline_sbom, current_sbom)

        # Typosquatting detection
        report.typosquatting_findings = self.typosquat_detector.scan_project(project_path)

        # License compliance (placeholder - would integrate with license_checker module)
        report.license_findings = self._check_licenses(project_path)

        # Dependency confusion (placeholder - would integrate with dep_confusion module)
        report.confusion_findings = self._check_dependency_confusion(project_path)

        # SBOM integrity check
        current_sbom = self._find_or_generate_sbom(project_path)
        if current_sbom:
            integrity = self.integrity_checker.check_integrity(current_sbom, project_path)
            report.integrity_issues = (
                integrity.undeclared_deps +
                integrity.phantom_components +
                integrity.checksum_mismatches
            )

        # Generate risk summary and recommendations
        report.risk_summary = self._collect_risk_summary(report)
        report.score = self._calculate_overall_score(report)
        report.recommendations = self._generate_recommendations(report)

        return report

    def _find_or_generate_sbom(self, project_path: str) -> Optional[str]:
        """Find existing SBOM or return None."""
        # Look for common SBOM file names
        sbom_patterns = ['sbom.json', 'bom.json', 'cyclonedx.json', 'spdx.json']
        for pattern in sbom_patterns:
            for sbom_file in Path(project_path).rglob(pattern):
                return str(sbom_file)
        return None

    def _check_licenses(self, project_path: str) -> List[Dict]:
        """Check license compliance (placeholder)."""
        # Would integrate with auditlens.license_checker if it exists
        return []

    def _check_dependency_confusion(self, project_path: str) -> List[Dict]:
        """Check for dependency confusion vulnerabilities (placeholder)."""
        # Would integrate with auditlens.dep_confusion if it exists
        return []

    def _collect_risk_summary(self, report: SupplyChainReport) -> Dict:
        """Collect risk summary statistics."""
        summary = {
            'total_components': 0,
            'high_risk': 0,
            'medium_risk': 0,
            'low_risk': 0,
            'typosquatting_detections': len(report.typosquatting_findings),
            'license_issues': len(report.license_findings),
            'confusion_vulnerabilities': len(report.confusion_findings),
            'integrity_issues': len(report.integrity_issues),
        }

        if report.sbom_diff:
            summary['total_components'] = (
                len(report.sbom_diff.added) +
                len(report.sbom_diff.removed) +
                len(report.sbom_diff.modified) +
                report.sbom_diff.unchanged
            )

            for change in report.sbom_diff.modified:
                if change.risk in ('CRITICAL', 'HIGH'):
                    summary['high_risk'] += 1
                elif change.risk == 'MEDIUM':
                    summary['medium_risk'] += 1
                else:
                    summary['low_risk'] += 1

        return summary

    def _calculate_overall_score(self, report: SupplyChainReport) -> float:
        """Calculate overall supply chain security score (0-100, higher is worse)."""
        score = 0.0

        # SBOM diff risk
        if report.sbom_diff:
            score += report.sbom_diff.risk_score * 0.3

        # Typosquatting findings
        score += len(report.typosquatting_findings) * 10

        # License issues
        score += len(report.license_findings) * 5

        # Dependency confusion
        score += len(report.confusion_findings) * 15

        # Integrity issues
        score += len(report.integrity_issues) * 8

        return min(score, 100.0)

    def _generate_recommendations(self, report: SupplyChainReport) -> List[str]:
        """Generate actionable recommendations based on findings."""
        recommendations = []

        if report.sbom_diff:
            if len(report.sbom_diff.added) > 5:
                recommendations.append(
                    f"Review {len(report.sbom_diff.added)} newly added dependencies for legitimacy and security"
                )

            high_risk_changes = [c for c in report.sbom_diff.modified if c.risk in ('HIGH', 'CRITICAL')]
            if high_risk_changes:
                recommendations.append(
                    f"Investigate {len(high_risk_changes)} high-risk version changes (downgrades detected)"
                )

        if report.typosquatting_findings:
            recommendations.append(
                f"Verify {len(report.typosquatting_findings)} packages flagged as potential typosquatting"
            )

        if report.confusion_findings:
            recommendations.append(
                "Implement namespace scoping for internal packages to prevent dependency confusion attacks"
            )

        if report.integrity_issues:
            recommendations.append(
                f"Update SBOM to reflect actual project dependencies ({len(report.integrity_issues)} discrepancies found)"
            )

        if not recommendations:
            recommendations.append("No critical supply chain security issues detected")

        return recommendations


# ── Report Generator ───────────────────────────────────────────────────────────

class SupplyChainReportGenerator:
    """Generates multi-format reports for supply chain audits."""

    def generate_html(self, report: SupplyChainReport, output_path: str) -> str:
        """Generate HTML report."""
        html = self._render_html_template(report)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)

        return output_path

    def generate_json(self, report: SupplyChainReport, output_path: str) -> str:
        """Generate JSON report."""
        # Convert dataclasses to dict
        report_dict = self._report_to_dict(report)

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report_dict, f, indent=2, default=str)

        return output_path

    def generate_xlsx(self, report: SupplyChainReport, output_path: str) -> str:
        """Generate XLSX report."""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment
        except ImportError:
            raise ImportError("openpyxl required for XLSX export: pip install openpyxl")

        wb = openpyxl.Workbook()

        # Summary sheet
        ws_summary = wb.active
        ws_summary.title = "Summary"
        self._write_summary_sheet(ws_summary, report)

        # SBOM Diff sheet
        if report.sbom_diff:
            ws_diff = wb.create_sheet("SBOM Changes")
            self._write_diff_sheet(ws_diff, report.sbom_diff)

        # Typosquatting sheet
        if report.typosquatting_findings:
            ws_typo = wb.create_sheet("Typosquatting")
            self._write_typosquatting_sheet(ws_typo, report.typosquatting_findings)

        wb.save(output_path)
        return output_path

    def generate_docx(self, report: SupplyChainReport, output_path: str) -> str:
        """Generate DOCX report."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
        except ImportError:
            raise ImportError("python-docx required for DOCX export: pip install python-docx")

        doc = Document()

        # Title
        title = doc.add_heading('Supply Chain Security Audit Report', 0)

        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(f"Generated: {report.generated_at}")
        doc.add_paragraph(f"Overall Risk Score: {report.score:.1f}/100")

        # Risk Summary
        doc.add_heading('Risk Summary', 2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Count'

        for key, value in report.risk_summary.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key.replace('_', ' ').title()
            row_cells[1].text = str(value)

        # Recommendations
        doc.add_heading('Recommendations', 1)
        for i, rec in enumerate(report.recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")

        # Details sections
        if report.sbom_diff:
            self._write_sbom_diff_section(doc, report.sbom_diff)

        if report.typosquatting_findings:
            self._write_typosquatting_section(doc, report.typosquatting_findings)

        doc.save(output_path)
        return output_path

    def _render_html_template(self, report: SupplyChainReport) -> str:
        """Render HTML report template."""
        html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Supply Chain Security Audit Report</title>
    <style>
        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; background: #f5f5f5; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: white; padding: 40px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        h1 {{ color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        .score {{ font-size: 48px; font-weight: bold; color: {self._risk_color(report.score)}; }}
        .summary-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 20px; margin: 20px 0; }}
        .summary-card {{ background: #ecf0f1; padding: 20px; border-radius: 8px; text-align: center; }}
        .summary-card .label {{ font-size: 14px; color: #7f8c8d; margin-bottom: 10px; }}
        .summary-card .value {{ font-size: 32px; font-weight: bold; color: #2c3e50; }}
        table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
        th, td {{ padding: 12px; text-align: left; border-bottom: 1px solid #ddd; }}
        th {{ background: #34495e; color: white; }}
        tr:hover {{ background: #f8f9fa; }}
        .risk-high {{ color: #e74c3c; font-weight: bold; }}
        .risk-medium {{ color: #f39c12; }}
        .risk-low {{ color: #27ae60; }}
        .recommendation {{ background: #e8f8f5; border-left: 4px solid #1abc9c; padding: 15px; margin: 10px 0; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Supply Chain Security Audit Report</h1>
        <p><strong>Generated:</strong> {report.generated_at}</p>

        <h2>Overall Risk Score</h2>
        <div class="score">{report.score:.1f}/100</div>

        <h2>Risk Summary</h2>
        <div class="summary-grid">
            {self._render_summary_cards(report.risk_summary)}
        </div>

        <h2>Recommendations</h2>
        {self._render_recommendations(report.recommendations)}

        {self._render_sbom_diff_html(report.sbom_diff) if report.sbom_diff else ''}
        {self._render_typosquatting_html(report.typosquatting_findings) if report.typosquatting_findings else ''}
        {self._render_integrity_issues_html(report.integrity_issues) if report.integrity_issues else ''}
    </div>
</body>
</html>
"""
        return html

    def _risk_color(self, score: float) -> str:
        """Get color for risk score."""
        if score >= 70:
            return '#e74c3c'
        elif score >= 40:
            return '#f39c12'
        else:
            return '#27ae60'

    def _render_summary_cards(self, summary: Dict) -> str:
        """Render summary cards HTML."""
        cards = []
        for key, value in summary.items():
            label = key.replace('_', ' ').title()
            cards.append(f"""
                <div class="summary-card">
                    <div class="label">{label}</div>
                    <div class="value">{value}</div>
                </div>
            """)
        return '\n'.join(cards)

    def _render_recommendations(self, recommendations: List[str]) -> str:
        """Render recommendations HTML."""
        items = [f'<div class="recommendation">{rec}</div>' for rec in recommendations]
        return '\n'.join(items)

    def _render_sbom_diff_html(self, diff: DiffResult) -> str:
        """Render SBOM diff section."""
        html = f"""
        <h2>SBOM Changes</h2>
        <p>Risk Score: <strong>{diff.risk_score:.1f}/100</strong></p>

        <h3>Added Components ({len(diff.added)})</h3>
        <table>
            <thead><tr><th>Name</th><th>Version</th><th>Ecosystem</th><th>License</th></tr></thead>
            <tbody>
                {self._render_component_rows(diff.added)}
            </tbody>
        </table>

        <h3>Modified Components ({len(diff.modified)})</h3>
        <table>
            <thead><tr><th>Name</th><th>Old Version</th><th>New Version</th><th>Change Type</th><th>Risk</th></tr></thead>
            <tbody>
                {self._render_change_rows(diff.modified)}
            </tbody>
        </table>

        <h3>Removed Components ({len(diff.removed)})</h3>
        <table>
            <thead><tr><th>Name</th><th>Version</th><th>Ecosystem</th></tr></thead>
            <tbody>
                {self._render_component_rows(diff.removed)}
            </tbody>
        </table>
        """
        return html

    def _render_component_rows(self, components: List[Component]) -> str:
        """Render component table rows."""
        rows = []
        for comp in components:
            rows.append(f"""
                <tr>
                    <td>{comp.name}</td>
                    <td>{comp.version}</td>
                    <td>{comp.ecosystem}</td>
                    <td>{comp.license}</td>
                </tr>
            """)
        return '\n'.join(rows) if rows else '<tr><td colspan="4">None</td></tr>'

    def _render_change_rows(self, changes: List[ComponentChange]) -> str:
        """Render change table rows."""
        rows = []
        for change in changes:
            risk_class = f"risk-{change.risk.lower()}"
            rows.append(f"""
                <tr>
                    <td>{change.name}</td>
                    <td>{change.old_version}</td>
                    <td>{change.new_version}</td>
                    <td>{change.change_type}</td>
                    <td class="{risk_class}">{change.risk}</td>
                </tr>
            """)
        return '\n'.join(rows) if rows else '<tr><td colspan="5">None</td></tr>'

    def _render_typosquatting_html(self, findings: List[TyposquattingMatch]) -> str:
        """Render typosquatting findings."""
        rows = []
        for match in findings:
            rows.append(f"""
                <tr>
                    <td class="risk-high">{match.suspicious_package}</td>
                    <td>{match.legitimate_package}</td>
                    <td>{match.similarity_score:.2f}</td>
                    <td>{match.detection_method}</td>
                    <td>{match.ecosystem}</td>
                </tr>
            """)

        html = f"""
        <h2>Typosquatting Detection ({len(findings)})</h2>
        <table>
            <thead><tr><th>Suspicious Package</th><th>Legitimate Package</th><th>Similarity</th><th>Method</th><th>Ecosystem</th></tr></thead>
            <tbody>
                {''.join(rows)}
            </tbody>
        </table>
        """
        return html

    def _render_integrity_issues_html(self, issues: List[Dict]) -> str:
        """Render integrity issues."""
        if not issues:
            return ''

        rows = []
        for issue in issues:
            rows.append(f"""
                <tr>
                    <td>{issue.get('name', 'N/A')}</td>
                    <td>{issue.get('severity', 'UNKNOWN')}</td>
                    <td>{issue.get('message', '')}</td>
                </tr>
            """)

        html = f"""
        <h2>SBOM Integrity Issues ({len(issues)})</h2>
        <table>
            <thead><tr><th>Component</th><th>Severity</th><th>Message</th></tr></thead>
            <tbody>
                {''.join(rows)}
            </tbody>
        </table>
        """
        return html

    def _report_to_dict(self, report: SupplyChainReport) -> Dict:
        """Convert report to dictionary."""
        def convert(obj):
            if hasattr(obj, '__dict__'):
                return {k: convert(v) for k, v in obj.__dict__.items()}
            elif isinstance(obj, list):
                return [convert(item) for item in obj]
            elif isinstance(obj, dict):
                return {k: convert(v) for k, v in obj.items()}
            else:
                return obj

        return convert(report)

    def _write_summary_sheet(self, ws, report):
        """Write summary to Excel sheet."""
        ws['A1'] = 'Supply Chain Security Audit'
        ws['A1'].font = Font(size=16, bold=True)

        ws['A3'] = 'Generated:'
        ws['B3'] = report.generated_at

        ws['A4'] = 'Overall Risk Score:'
        ws['B4'] = f"{report.score:.1f}/100"

        row = 6
        ws[f'A{row}'] = 'Risk Summary'
        ws[f'A{row}'].font = Font(bold=True)
        row += 1

        for key, value in report.risk_summary.items():
            ws[f'A{row}'] = key.replace('_', ' ').title()
            ws[f'B{row}'] = value
            row += 1

    def _write_diff_sheet(self, ws, diff):
        """Write SBOM diff to Excel sheet."""
        ws['A1'] = 'SBOM Changes'
        ws['A1'].font = Font(size=14, bold=True)

        row = 3
        ws[f'A{row}'] = 'Added Components'
        ws[f'A{row}'].font = Font(bold=True)
        row += 1

        headers = ['Name', 'Version', 'Ecosystem', 'License']
        for col, header in enumerate(headers, 1):
            ws.cell(row, col, header).font = Font(bold=True)
        row += 1

        for comp in diff.added:
            ws.cell(row, 1, comp.name)
            ws.cell(row, 2, comp.version)
            ws.cell(row, 3, comp.ecosystem)
            ws.cell(row, 4, comp.license)
            row += 1

    def _write_typosquatting_sheet(self, ws, findings):
        """Write typosquatting findings to Excel sheet."""
        ws['A1'] = 'Typosquatting Detection'
        ws['A1'].font = Font(size=14, bold=True)

        headers = ['Suspicious Package', 'Legitimate Package', 'Similarity', 'Method', 'Ecosystem']
        for col, header in enumerate(headers, 1):
            ws.cell(3, col, header).font = Font(bold=True)

        for row, match in enumerate(findings, 4):
            ws.cell(row, 1, match.suspicious_package)
            ws.cell(row, 2, match.legitimate_package)
            ws.cell(row, 3, match.similarity_score)
            ws.cell(row, 4, match.detection_method)
            ws.cell(row, 5, match.ecosystem)

    def _write_sbom_diff_section(self, doc, diff):
        """Write SBOM diff section to Word document."""
        doc.add_heading('SBOM Changes', 1)
        doc.add_paragraph(f"Risk Score: {diff.risk_score:.1f}/100")

        doc.add_heading('Added Components', 2)
        if diff.added:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Version'
            hdr_cells[2].text = 'Ecosystem'
            hdr_cells[3].text = 'License'

            for comp in diff.added[:50]:  # Limit to prevent large files
                row_cells = table.add_row().cells
                row_cells[0].text = comp.name
                row_cells[1].text = comp.version
                row_cells[2].text = comp.ecosystem
                row_cells[3].text = comp.license
        else:
            doc.add_paragraph('None')

    def _write_typosquatting_section(self, doc, findings):
        """Write typosquatting section to Word document."""
        doc.add_heading('Typosquatting Detection', 1)

        if findings:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Suspicious'
            hdr_cells[1].text = 'Legitimate'
            hdr_cells[2].text = 'Similarity'
            hdr_cells[3].text = 'Method'
            hdr_cells[4].text = 'Ecosystem'

            for match in findings[:50]:
                row_cells = table.add_row().cells
                row_cells[0].text = match.suspicious_package
                row_cells[1].text = match.legitimate_package
                row_cells[2].text = f"{match.similarity_score:.2f}"
                row_cells[3].text = match.detection_method
                row_cells[4].text = match.ecosystem
        else:
            doc.add_paragraph('No typosquatting detected')


# ── Public API Functions ───────────────────────────────────────────────────────

def diff_sboms(baseline_sbom: str, current_sbom: str, format: str = 'cyclonedx') -> DiffResult:
    """Compare two SBOMs and return differences."""
    differ = SBOMDiffer()
    return differ.diff(baseline_sbom, current_sbom, format)


def detect_supply_chain_attacks(project_path: str, baseline_path: Optional[str] = None) -> List[dict]:
    """Detect supply chain attacks: dependency confusion, typosquatting, etc."""
    guard = SupplyChainGuard()
    report = guard.audit_project(project_path, baseline_path)

    attacks = []

    # Typosquatting
    for finding in report.typosquatting_findings:
        attacks.append({
            'type': 'typosquatting',
            'severity': 'HIGH',
            'package': finding.suspicious_package,
            'message': f"Potential typosquatting: {finding.suspicious_package} similar to {finding.legitimate_package}",
            'details': asdict(finding)
        })

    # Suspicious SBOM changes
    if report.sbom_diff:
        for change in report.sbom_diff.modified:
            if change.risk in ('HIGH', 'CRITICAL'):
                attacks.append({
                    'type': 'suspicious_version_change',
                    'severity': change.risk,
                    'package': change.name,
                    'message': change.reason,
                    'details': asdict(change)
                })

    return attacks


def audit_supply_chain(project_path: str, severity: str = 'LOW') -> SupplyChainReport:
    """Execute complete supply chain audit."""
    guard = SupplyChainGuard()
    return guard.audit_project(project_path, None, severity)


def generate_supply_chain_report(
    report: SupplyChainReport,
    format: str = 'html',
    output_path: Optional[str] = None
) -> str:
    """Generate supply chain audit report in specified format."""
    generator = SupplyChainReportGenerator()

    if output_path is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = f'supply_chain_report_{timestamp}.{format}'

    if format == 'html':
        return generator.generate_html(report, output_path)
    elif format == 'json':
        return generator.generate_json(report, output_path)
    elif format == 'xlsx':
        return generator.generate_xlsx(report, output_path)
    elif format == 'docx':
        return generator.generate_docx(report, output_path)
    else:
        raise ValueError(f"Unsupported format: {format}. Use 'html', 'json', 'xlsx', or 'docx'")


def detect_typosquatting(project_path: str, threshold: float = 0.85) -> List[dict]:
    """Detect potential typosquatting packages."""
    detector = TyposquattingDetector(threshold)
    matches = detector.scan_project(project_path)

    return [
        {
            'suspicious_package': m.suspicious_package,
            'legitimate_package': m.legitimate_package,
            'similarity_score': m.similarity_score,
            'detection_method': m.detection_method,
            'ecosystem': m.ecosystem,
            'severity': 'HIGH',
            'message': f"Package '{m.suspicious_package}' is suspiciously similar to popular package '{m.legitimate_package}'"
        }
        for m in matches
    ]


def check_sbom_integrity(sbom_path: str, project_path: str) -> IntegrityReport:
    """Verify SBOM integrity against actual project dependencies."""
    checker = SBOMIntegrityChecker()
    return checker.check_integrity(sbom_path, project_path)
