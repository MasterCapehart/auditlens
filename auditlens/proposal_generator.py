"""
AuditLens — Generador de Propuesta Comercial + Contrato de Auditoría

Genera un documento Word profesional listo para enviar al cliente con:
- Alcance de la auditoría
- Metodología
- Precio estimado automático según tamaño del proyecto
- Timeline de trabajo
- Contrato de servicios tipo
"""
from __future__ import annotations

import os
from datetime import date, timedelta
from typing import Any, Dict, List, Optional


def _estimate_price(project_path: str, frameworks: List[str]) -> Dict[str, Any]:
    """Estimate audit price based on project size and frameworks requested."""
    try:
        total_files = sum(
            len(files)
            for _, _, files in os.walk(project_path)
            if not any(skip in _ for skip in ('venv', 'node_modules', '.git', '__pycache__'))
        )
    except Exception:
        total_files = 50

    # Base price tiers (USD)
    if total_files < 30:
        base = 800_000      # CLP
        size_label = 'Pequeño (<30 archivos)'
        days = 3
    elif total_files < 100:
        base = 1_500_000
        size_label = 'Mediano (30-100 archivos)'
        days = 5
    elif total_files < 300:
        base = 2_800_000
        size_label = 'Grande (100-300 archivos)'
        days = 10
    else:
        base = 5_000_000
        size_label = 'Muy grande (300+ archivos)'
        days = 15

    # Per-framework addon
    fw_prices = {
        'ley21719':  300_000,
        'iso27001':  400_000,
        'cmf':       500_000,
        'gdpr':      350_000,
        'hipaa':     350_000,
        'pci':       400_000,
    }
    fw_total = sum(fw_prices.get(fw, 200_000) for fw in frameworks)
    total = base + fw_total

    return {
        'base': base,
        'frameworks_addon': fw_total,
        'total': total,
        'total_uf': round(total / 38_500, 1),   # approx UF value
        'size_label': size_label,
        'estimated_days': days,
        'files': total_files,
    }


def generate_proposal_docx(
    project_path: str,
    output_path: str,
    empresa_cliente: str = 'Empresa Cliente',
    rut_cliente: str = '',
    contacto_cliente: str = '',
    empresa_auditora: str = 'AuditLens Security',
    auditor: str = '',
    rut_auditor: str = '',
    frameworks: Optional[List[str]] = None,
    incluir_contrato: bool = True,
) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print('\033[91m[AuditLens]\033[0m python-docx no instalado: pip install python-docx')
        return ''

    frameworks = frameworks or ['sast', 'sca', 'ley21719']
    today      = date.today()
    today_str  = today.strftime('%d de %B de %Y').replace(
        'January','enero').replace('February','febrero').replace('March','marzo').replace(
        'April','abril').replace('May','mayo').replace('June','junio').replace(
        'July','julio').replace('August','agosto').replace('September','septiembre').replace(
        'October','octubre').replace('November','noviembre').replace('December','diciembre')
    valid_until = (today + timedelta(days=30)).strftime('%d/%m/%Y')
    pricing     = _estimate_price(project_path, frameworks)
    start_date  = (today + timedelta(days=7)).strftime('%d/%m/%Y')
    end_date    = (today + timedelta(days=7 + pricing['estimated_days'])).strftime('%d/%m/%Y')

    fw_label_map = {
        'sast': 'Análisis estático de código (SAST)',
        'sca': 'Análisis de composición de software (SCA)',
        'ley21719': 'Cumplimiento Ley 21.719 — Protección de Datos Personales',
        'iso27001': 'Cumplimiento ISO 27001:2022',
        'cmf': 'Cumplimiento CMF Circular 57 / Norma 461',
        'gdpr': 'Cumplimiento GDPR',
        'hipaa': 'Cumplimiento HIPAA',
        'pci': 'Cumplimiento PCI-DSS v4.0',
        'dast': 'Análisis dinámico de seguridad (DAST)',
        'git': 'Análisis de historial Git (secretos, arqueología)',
    }

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    def heading(text, level=1):
        p = doc.add_heading(text, level=level)
        if p.runs:
            p.runs[0].font.color.rgb = RGBColor(0x1f, 0x6f, 0xeb)
        return p

    def para(text='', bold=False, italic=False, size=11, color=None):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        return p

    def table2(rows):
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = 'Table Grid'
        for i, (a, b) in enumerate(rows):
            t.rows[i].cells[0].text = a
            t.rows[i].cells[1].text = b
            t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            for cell in t.rows[i].cells:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
        return t

    # ── Portada ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('PROPUESTA DE SERVICIOS\nAUDITORÍA DE CIBERSEGURIDAD')
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1f, 0x6f, 0xeb)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Preparado para: {empresa_cliente}\n').bold = True
    if rut_cliente:
        meta.add_run(f'RUT: {rut_cliente}\n')
    meta.add_run(f'Fecha: {today_str}\n')
    meta.add_run(f'Válida hasta: {valid_until}\n')
    meta.add_run(f'Preparado por: {empresa_auditora}')
    if auditor:
        meta.add_run(f' — {auditor}')
    doc.add_page_break()

    # ── 1. Resumen ejecutivo ──────────────────────────────────────────────────
    heading('1. Resumen Ejecutivo')
    para(
        f'{empresa_auditora} propone realizar una auditoría integral de ciberseguridad '
        f'para {empresa_cliente}. El objetivo es identificar vulnerabilidades técnicas, '
        f'evaluar el cumplimiento normativo y entregar un plan de remediación priorizado '
        f'con recomendaciones concretas.'
    )

    # ── 2. Alcance ────────────────────────────────────────────────────────────
    heading('2. Alcance del Servicio')
    para('El servicio incluye los siguientes módulos de análisis:', bold=True)
    for fw in frameworks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(fw_label_map.get(fw, fw.upper()))

    doc.add_paragraph()
    para('Tamaño estimado del proyecto auditado:', bold=True)
    table2([
        ('Archivos a analizar', str(pricing['files'])),
        ('Clasificación', pricing['size_label']),
        ('Días de trabajo estimados', str(pricing['estimated_days'])),
    ])

    # ── 3. Metodología ────────────────────────────────────────────────────────
    heading('3. Metodología')
    steps = [
        ('Fase 1 — Reconocimiento', 'Análisis del código fuente, dependencias, configuraciones e infraestructura.'),
        ('Fase 2 — Análisis estático', 'SAST automatizado con reglas propias + análisis manual de hallazgos críticos.'),
        ('Fase 3 — Análisis de composición', 'SCA: CVEs en dependencias, licencias, dependency confusion.'),
        ('Fase 4 — Compliance', 'Mapeo de hallazgos a frameworks regulatorios aplicables.'),
        ('Fase 5 — Reporte', 'Entrega de informe técnico + informe ejecutivo + plan de remediación.'),
        ('Fase 6 — Presentación', 'Sesión de presentación de resultados (1h, videoconferencia).'),
    ]
    for fase, desc in steps:
        p = doc.add_paragraph()
        p.add_run(f'{fase}: ').bold = True
        p.add_run(desc)

    # ── 4. Entregables ────────────────────────────────────────────────────────
    heading('4. Entregables')
    deliverables = [
        'Informe técnico completo (PDF + Word) con todos los hallazgos',
        'Informe ejecutivo para directorio (sin jerga técnica)',
        'Reporte de compliance por framework (OWASP, ISO 27001, Ley 21.719, etc.)',
        'Plan de remediación priorizado (CRÍTICO → BAJO)',
        'Archivo SARIF para integración con herramientas DevSecOps',
        'Sesión de presentación de resultados (1h videoconferencia)',
        'Soporte post-auditoría por 30 días para consultas',
    ]
    for d in deliverables:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(d)

    # ── 5. Timeline ───────────────────────────────────────────────────────────
    heading('5. Cronograma')
    table2([
        ('Inicio estimado', start_date),
        ('Fin estimado', end_date),
        ('Duración', f'{pricing["estimated_days"]} días hábiles'),
        ('Entrega del informe', f'{pricing["estimated_days"]} días desde el inicio'),
        ('Presentación de resultados', f'{pricing["estimated_days"] + 2} días desde el inicio'),
    ])

    # ── 6. Inversión ──────────────────────────────────────────────────────────
    heading('6. Inversión')
    table2([
        ('Análisis base (SAST + SCA + Informe)',
         f'${pricing["base"]:,.0f} CLP'),
        ('Módulos de compliance adicionales',
         f'${pricing["frameworks_addon"]:,.0f} CLP'),
        ('TOTAL', f'${pricing["total"]:,.0f} CLP ({pricing["total_uf"]} UF aprox.)'),
        ('Condiciones de pago', '50% al inicio · 50% a la entrega del informe'),
        ('Validez de la propuesta', valid_until),
    ])
    para('* Los precios no incluyen IVA (19%).',
         italic=True, size=9, color=RGBColor(0x8b, 0x94, 0x9e))

    # ── 7. Sobre nosotros ─────────────────────────────────────────────────────
    heading('7. Sobre Nosotros')
    para(
        f'{empresa_auditora} es una empresa especializada en ciberseguridad y auditoría '
        f'de sistemas de información. Utilizamos herramientas propias de análisis automatizado '
        f'combinadas con revisión manual por expertos certificados. Nos especializamos en el '
        f'marco regulatorio chileno (Ley 21.719, CMF, Ley de Ciberseguridad).'
    )

    if incluir_contrato:
        doc.add_page_break()
        # ── Contrato ──────────────────────────────────────────────────────────
        heading('CONTRATO DE SERVICIOS DE AUDITORÍA DE CIBERSEGURIDAD')
        para(
            f'En Santiago, a {today_str}, entre {empresa_auditora} (en adelante "el Prestador"), '
            f'representado por {auditor or "su representante legal"}'
            + (f', RUT {rut_auditor}' if rut_auditor else '') +
            f'; y {empresa_cliente} (en adelante "el Cliente")'
            + (f', RUT {rut_cliente}' if rut_cliente else '') +
            f', se celebra el siguiente contrato:'
        )
        clauses = [
            ('Cláusula 1ª — Objeto',
             'El Prestador se obliga a realizar una auditoría de ciberseguridad del sistema '
             'de información indicado por el Cliente, conforme al alcance definido en la propuesta '
             'adjunta, que forma parte integrante de este contrato.'),
            ('Cláusula 2ª — Precio y Forma de Pago',
             f'El precio total del servicio es de ${pricing["total"]:,.0f} CLP (sin IVA). '
             f'Se pagará 50% al inicio de los trabajos y 50% a la entrega del informe final.'),
            ('Cláusula 3ª — Confidencialidad',
             'El Prestador se obliga a mantener estricta confidencialidad sobre toda la información '
             'a la que acceda en el marco de este contrato, y no divulgará hallazgos a terceros '
             'sin autorización escrita del Cliente.'),
            ('Cláusula 4ª — Propiedad Intelectual',
             'Los informes y documentos generados son de propiedad del Cliente una vez efectuado '
             'el pago total. Las herramientas y metodologías utilizadas son de propiedad del Prestador.'),
            ('Cláusula 5ª — Limitación de Responsabilidad',
             'El Prestador responde por la calidad técnica del servicio, pero no garantiza la '
             'ausencia total de vulnerabilidades. La auditoría refleja el estado del sistema a '
             'la fecha del análisis.'),
            ('Cláusula 6ª — Duración',
             f'Este contrato tiene una duración de {pricing["estimated_days"] + 32} días corridos '
             f'desde su firma, incluyendo el período de soporte post-entrega.'),
        ]
        for title, text in clauses:
            p = doc.add_paragraph()
            p.add_run(f'{title}: ').bold = True
            p.add_run(text)

        doc.add_paragraph()
        doc.add_paragraph()
        sig = doc.add_table(rows=2, cols=2)
        sig.rows[0].cells[0].text = 'El Prestador'
        sig.rows[0].cells[1].text = 'El Cliente'
        sig.rows[1].cells[0].text = f'{auditor or "_______________"}\n{empresa_auditora}'
        sig.rows[1].cells[1].text = f'{contacto_cliente or "_______________"}\n{empresa_cliente}'

    doc.save(output_path)
    print(f'\033[92m[AuditLens]\033[0m Propuesta generada: {output_path}')
    return output_path
