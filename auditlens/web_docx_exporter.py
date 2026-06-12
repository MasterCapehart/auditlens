"""
AuditLens Web DOCX Exporter — informe Word de auditoría de seguridad web.

Genera un documento profesional con:
  - Portada
  - Tabla de contenidos
  - 1. Resumen Ejecutivo
  - 2. Alcance y Metodología (OWASP Testing Guide v4.2)
  - 3. Inventario Técnico (tech stack, páginas crawleadas, SSL info)
  - 4. Hallazgos por Categoría (SSL, Headers, Cookies, CORS, Forms, XSS, SQLi, Paths, JS)
  - 5. Análisis de Brechas OWASP Top 10 : 2021
  - 6. Mapa de Riesgo
  - 7. Conclusiones
  - 8. Recomendaciones Priorizadas
  - 9. Plan de Remediación con SLAs
  - 10. Anexos (payloads usados, herramientas, glosario)

Requires: pip install python-docx
"""

from __future__ import annotations

import os
from datetime import datetime
from typing import Dict, List, Optional

_SEVERITY_ORDER = {'CRITICAL': 0, 'HIGH': 1, 'MEDIUM': 2, 'LOW': 3}

_SEV_COLORS = {
    'CRITICAL': (220, 0, 0),
    'HIGH': (255, 102, 0),
    'MEDIUM': (204, 153, 0),
    'LOW': (0, 102, 204),
}

_OWASP_CATEGORIES = {
    'A01': 'Broken Access Control',
    'A02': 'Cryptographic Failures',
    'A03': 'Injection',
    'A04': 'Insecure Design',
    'A05': 'Security Misconfiguration',
    'A06': 'Vulnerable & Outdated Components',
    'A07': 'Identification & Authentication Failures',
    'A08': 'Software & Data Integrity Failures',
    'A09': 'Security Logging & Monitoring Failures',
    'A10': 'Server-Side Request Forgery',
}

_CATEGORY_DESCRIPTIONS = {
    'SSL/TLS': 'Evaluación del protocolo SSL/TLS, versiones soportadas, ciphers y validez del certificado.',
    'HTTP Headers': 'Análisis de cabeceras de seguridad HTTP que protegen contra XSS, clickjacking y sniffing.',
    'Cookies': 'Revisión de atributos de seguridad en cookies (Secure, HttpOnly, SameSite).',
    'CORS': 'Validación de la política de Cross-Origin Resource Sharing.',
    'Forms': 'Análisis de formularios HTML: CSRF tokens, action HTTPS, autocomplete.',
    'JavaScript': 'Análisis estático (SAST) de código JavaScript descargado desde la aplicación.',
    'Sensitive Paths': 'Enumeración de rutas sensibles accesibles públicamente.',
    'XSS': 'Pruebas de Cross-Site Scripting reflejado en parámetros GET/POST.',
    'SQLi': 'Detección de mensajes de error que revelan vulnerabilidades de SQL Injection.',
    'Open Redirect': 'Detección de parámetros que permiten redirecciones abiertas a sitios externos.',
    'Info Disclosure': 'Identificación de información sensible expuesta en respuestas HTTP.',
}

_RULE_CATEGORY_MAP = {
    'WEB-SSL': 'SSL/TLS',
    'WEB-HDR': 'HTTP Headers',
    'WEB-CKI': 'Cookies',
    'WEB-CRS': 'CORS',
    'WEB-FRM': 'Forms',
    'WEB-XSS': 'XSS',
    'WEB-SQL': 'SQLi',
    'WEB-RDR': 'Open Redirect',
    'WEB-PTH': 'Sensitive Paths',
    'WEB-INF': 'Info Disclosure',
    'SEC-': 'JavaScript',
    'AST-': 'JavaScript',
    'TAINT-': 'JavaScript',
}

_SLA_BY_SEVERITY = {
    'CRITICAL': '72 horas',
    'HIGH': '7 días',
    'MEDIUM': '30 días',
    'LOW': '90 días',
}

_REMEDIATION_GUIDES = {
    'SSL/TLS': (
        'Configurar el servidor para soportar únicamente TLS 1.2 y TLS 1.3. '
        'Desactivar SSLv3, TLS 1.0 y TLS 1.1. Usar ciphers ECDHE+AES256+GCM. '
        'Renovar el certificado con antelación mínima de 30 días.'
    ),
    'HTTP Headers': (
        'Agregar en la configuración del servidor web (nginx/Apache/IIS): '
        'Strict-Transport-Security, Content-Security-Policy, X-Frame-Options: DENY, '
        'X-Content-Type-Options: nosniff, Referrer-Policy: strict-origin-when-cross-origin, '
        'Permissions-Policy. Revisar y restringir Server y X-Powered-By.'
    ),
    'Cookies': (
        'En todas las cookies de sesión agregar los atributos: Secure; HttpOnly; SameSite=Strict. '
        'Verificar que el framework establezca estos atributos automáticamente en las cookies de sesión.'
    ),
    'CORS': (
        'Reemplazar Access-Control-Allow-Origin: * por una whitelist explícita de orígenes autorizados. '
        'Nunca combinar Access-Control-Allow-Credentials: true con origen wildcard.'
    ),
    'Forms': (
        'Implementar el patrón Synchronized Token Pattern (CSRF token) en todos los formularios POST. '
        'Asegurarse de que todos los action de formularios usen HTTPS. '
        'Agregar autocomplete=off en campos de contraseña.'
    ),
    'JavaScript': (
        'Revisar el código JavaScript para eliminar credenciales hardcodeadas, endpoints internos expuestos '
        'y uso de algoritmos de cifrado débiles. Implementar Content-Security-Policy para restringir '
        'la ejecución de scripts inline.'
    ),
    'Sensitive Paths': (
        'Mover archivos de configuración fuera del DocumentRoot. Configurar reglas deny en el servidor web '
        'para .env, .git, .svn y rutas de administración. Implementar autenticación multifactor en /admin.'
    ),
    'XSS': (
        'Codificar HTML-encode todos los datos de usuario antes de insertarlos en respuestas. '
        'Implementar una Content-Security-Policy estricta. Usar frameworks que escapen por defecto (React, Angular, Vue).'
    ),
    'SQLi': (
        'Usar consultas parametrizadas (prepared statements) en toda la capa de acceso a datos. '
        'Nunca concatenar input de usuario en queries SQL. Configurar mensajes de error genéricos en producción.'
    ),
    'Open Redirect': (
        'Crear una whitelist de dominios permitidos para redirecciones. '
        'Validar y sanitizar todos los parámetros que controlen redirecciones. '
        'Usar redirecciones internas en lugar de aceptar URLs externas como parámetro.'
    ),
    'Info Disclosure': (
        'Desactivar debug mode en producción (DEBUG=False). Configurar páginas de error personalizadas. '
        'Suprimir o genericizar los headers Server y X-Powered-By. '
        'Configurar logging interno sin exponer detalles en las respuestas HTTP.'
    ),
}


def _get_category(rule_id: str) -> str:
    for prefix, cat in _RULE_CATEGORY_MAP.items():
        if rule_id.startswith(prefix):
            return cat
    return 'Otro'


def generate_web_docx_report(
    scan_result,
    output_path: str = 'informe_auditoria_web.docx',
    empresa: str = 'Empresa',
    sistema: str = 'Aplicación Web',
    auditor: str = '[Auditor por asignar]',
    authorized_by: str = '[Responsable técnico]',
) -> str:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError(
            'python-docx requerido: pip install python-docx --break-system-packages'
        )

    findings = scan_result.findings
    target_url = scan_result.target_url
    fecha = datetime.now().strftime('%d/%m/%Y')
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M')

    counts = {'CRITICAL': 0, 'HIGH': 0, 'MEDIUM': 0, 'LOW': 0}
    for f in findings:
        sev = f.get('severity', 'LOW').upper()
        if sev in counts:
            counts[sev] += 1

    doc = Document()

    # ── Estilos base ──────────────────────────────────────────────────────────
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    def _rgb(r, g, b):
        return RGBColor(r, g, b)

    def _h1(text: str):
        p = doc.add_heading(text, level=1)
        p.runs[0].font.color.rgb = _rgb(30, 58, 95)
        return p

    def _h2(text: str):
        p = doc.add_heading(text, level=2)
        p.runs[0].font.color.rgb = _rgb(30, 58, 95)
        return p

    def _h3(text: str):
        p = doc.add_heading(text, level=3)
        p.runs[0].font.color.rgb = _rgb(71, 85, 105)
        return p

    def _para(text: str, bold: bool = False, italic: bool = False, size: int = 11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        return p

    def _sev_color(sev: str):
        r, g, b = _SEV_COLORS.get(sev.upper(), (100, 116, 139))
        return _rgb(r, g, b)

    def _table_header_row(table, headers: list):
        row = table.rows[0]
        for i, h in enumerate(headers):
            cell = row.cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = _rgb(255, 255, 255)
            run.font.size = Pt(10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1E3A5F')
            tcPr.append(shd)

    def _set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    _SEV_HEX = {
        'CRITICAL': 'DC0000',
        'HIGH': 'FF6600',
        'MEDIUM': 'CC9900',
        'LOW': '0066CC',
    }

    # ─────────────────────────────────────────────────────────────────────────
    # PORTADA
    # ─────────────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('INFORME DE AUDITORÍA DE SEGURIDAD WEB')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _rgb(30, 58, 95)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(empresa)
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(sistema)
    run.font.size = Pt(14)
    run.font.color.rgb = _rgb(71, 85, 105)

    doc.add_paragraph()
    doc.add_paragraph()

    cover_table = doc.add_table(rows=7, cols=2)
    cover_table.style = 'Table Grid'
    cover_data = [
        ('URL auditada', target_url),
        ('Empresa', empresa),
        ('Sistema', sistema),
        ('Auditor', auditor),
        ('Autorizado por', authorized_by),
        ('Fecha', fecha),
        ('Clasificación', 'CONFIDENCIAL — Solo para uso interno'),
    ]
    for i, (k, v) in enumerate(cover_data):
        cover_table.rows[i].cells[0].text = k
        cover_table.rows[i].cells[1].text = v
        cover_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_bg(cover_table.rows[i].cells[0], 'E8EFF7')

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # TABLA DE CONTENIDOS (manual)
    # ─────────────────────────────────────────────────────────────────────────
    _h1('Tabla de Contenidos')
    toc_items = [
        '1. Resumen Ejecutivo',
        '2. Alcance y Metodología',
        '3. Inventario Técnico',
        '4. Hallazgos por Categoría',
        '5. Análisis de Brechas OWASP Top 10 : 2021',
        '6. Mapa de Riesgo',
        '7. Conclusiones',
        '8. Recomendaciones Priorizadas',
        '9. Plan de Remediación con SLAs',
        '10. Anexos',
    ]
    for item in toc_items:
        _para(item)
    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 1. RESUMEN EJECUTIVO
    # ─────────────────────────────────────────────────────────────────────────
    _h1('1. Resumen Ejecutivo')

    total = len(findings)
    risk_level = 'CRÍTICO' if counts['CRITICAL'] > 0 else ('ALTO' if counts['HIGH'] > 0 else 'MEDIO')
    risk_color = _rgb(220, 0, 0) if risk_level == 'CRÍTICO' else (_rgb(255, 102, 0) if risk_level == 'ALTO' else _rgb(204, 153, 0))

    p = doc.add_paragraph()
    p.add_run('Nivel de Riesgo Global: ').bold = True
    run = p.add_run(risk_level)
    run.bold = True
    run.font.color.rgb = risk_color

    _para(
        f'Se realizó una auditoría de seguridad web sobre {target_url} el {now_str}. '
        f'El escaneo analizó {scan_result.pages_scanned} páginas, '
        f'{scan_result.js_files_scanned} archivos JavaScript y '
        f'{scan_result.forms_scanned} formularios HTML, con una duración total de '
        f'{scan_result.scan_duration:.0f} segundos.'
    )
    _para(
        f'Se identificaron {total} hallazgos de seguridad: '
        f'{counts["CRITICAL"]} CRÍTICOS, {counts["HIGH"]} ALTOS, '
        f'{counts["MEDIUM"]} MEDIOS y {counts["LOW"]} BAJOS. '
        'Los hallazgos críticos y altos requieren atención inmediata.'
    )

    # Summary table
    doc.add_paragraph()
    sum_table = doc.add_table(rows=5, cols=3)
    sum_table.style = 'Table Grid'
    _table_header_row(sum_table, ['Severidad', 'Cantidad', 'SLA de Remediación'])
    rows_data = [
        ('CRITICAL', str(counts['CRITICAL']), _SLA_BY_SEVERITY['CRITICAL']),
        ('HIGH', str(counts['HIGH']), _SLA_BY_SEVERITY['HIGH']),
        ('MEDIUM', str(counts['MEDIUM']), _SLA_BY_SEVERITY['MEDIUM']),
        ('LOW', str(counts['LOW']), _SLA_BY_SEVERITY['LOW']),
    ]
    for i, (sev, cnt, sla) in enumerate(rows_data, 1):
        row = sum_table.rows[i]
        row.cells[0].text = sev
        row.cells[1].text = cnt
        row.cells[2].text = sla
        _set_cell_bg(row.cells[0], _SEV_HEX[sev])
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = _rgb(255, 255, 255)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 2. ALCANCE Y METODOLOGÍA
    # ─────────────────────────────────────────────────────────────────────────
    _h1('2. Alcance y Metodología')

    _h2('2.1 Alcance')
    _para(f'La auditoría cubre la aplicación web alojada en {target_url} y sus recursos '
          'accesibles públicamente, incluyendo subdirectorios, formularios, archivos JavaScript '
          'y cabeceras HTTP. No se auditaron sistemas de backend, bases de datos ni redes internas.')

    _para('Módulos auditados:', bold=True)
    modules = [
        '• SSL/TLS: versión de protocolo, cipher suites, expiración de certificado',
        '• Cabeceras HTTP de seguridad: CSP, HSTS, X-Frame-Options, CORS',
        '• Cookies: atributos Secure, HttpOnly, SameSite',
        '• Formularios HTML: CSRF tokens, acción HTTPS, autocomplete',
        '• JavaScript: análisis estático SAST sobre scripts descargados',
        '• Rutas sensibles: enumeración de recursos expuestos',
        '• DAST: XSS reflejado, SQLi error-based, open redirect',
        '• Divulgación de información: stack traces, versiones, directory listing',
    ]
    for m in modules:
        _para(m)

    _h2('2.2 Metodología')
    _para(
        'La metodología aplicada sigue el OWASP Testing Guide v4.2 (OTG) y el '
        'OWASP Web Security Testing Guide (WSTG), complementada con controles '
        'del OWASP Top 10 : 2021. Para cada categoría de prueba se aplicaron:'
    )
    method_steps = [
        '1. Reconocimiento pasivo — análisis de cabeceras, certificados y tecnología',
        '2. Enumeración activa — crawling de páginas y rutas sensibles',
        '3. Pruebas funcionales — inyección de payloads en parámetros controlables',
        '4. Análisis estático — revisión del código JavaScript disponible',
        '5. Validación — verificación manual de hallazgos positivos',
    ]
    for s in method_steps:
        _para(s)

    _h2('2.3 Herramientas')
    tools_table = doc.add_table(rows=6, cols=2)
    tools_table.style = 'Table Grid'
    _table_header_row(tools_table, ['Herramienta', 'Propósito'])
    tools = [
        ('AuditLens WebScanner', 'Motor principal — SSL, headers, cookies, CORS, DAST'),
        ('requests (Python)', 'HTTP client para crawling y probes'),
        ('ssl (Python stdlib)', 'Inspección de certificados y ciphers TLS'),
        ('Reglas SAST AuditLens', 'Análisis estático de JavaScript descargado'),
        ('OWASP WSTG v4.2', 'Marco metodológico de referencia'),
    ]
    for i, (tool, purpose) in enumerate(tools, 1):
        tools_table.rows[i].cells[0].text = tool
        tools_table.rows[i].cells[1].text = purpose

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 3. INVENTARIO TÉCNICO
    # ─────────────────────────────────────────────────────────────────────────
    _h1('3. Inventario Técnico')

    tech_table = doc.add_table(rows=6, cols=2)
    tech_table.style = 'Table Grid'
    _table_header_row(tech_table, ['Elemento', 'Valor'])
    tech_data = [
        ('URL objetivo', target_url),
        ('Páginas escaneadas', str(scan_result.pages_scanned)),
        ('Archivos JavaScript analizados', str(scan_result.js_files_scanned)),
        ('Formularios analizados', str(scan_result.forms_scanned)),
        ('Duración del escaneo', f'{scan_result.scan_duration:.1f} segundos'),
    ]
    for i, (k, v) in enumerate(tech_data, 1):
        tech_table.rows[i].cells[0].text = k
        tech_table.rows[i].cells[1].text = v
        tech_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_bg(tech_table.rows[i].cells[0], 'E8EFF7')

    if scan_result.tech_stack:
        _h2('3.1 Tecnología Detectada')
        for t in scan_result.tech_stack:
            _para(f'• {t}')

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 4. HALLAZGOS POR CATEGORÍA
    # ─────────────────────────────────────────────────────────────────────────
    _h1('4. Hallazgos por Categoría')

    # Group by category
    by_category: Dict[str, List[dict]] = {}
    for f in findings:
        cat = _get_category(f['rule_id'])
        by_category.setdefault(cat, []).append(f)

    section_num = 1
    for cat, cat_findings in sorted(by_category.items()):
        _h2(f'4.{section_num} {cat}')
        desc = _CATEGORY_DESCRIPTIONS.get(cat, '')
        if desc:
            _para(desc, italic=True)
        _para(f'Total hallazgos en esta categoría: {len(cat_findings)}')

        # Table of findings
        t = doc.add_table(rows=1 + len(cat_findings), cols=5)
        t.style = 'Table Grid'
        _table_header_row(t, ['Severidad', 'ID Regla', 'Nombre', 'URL / Archivo', 'Compliance'])

        for j, f in enumerate(sorted(cat_findings, key=lambda x: _SEVERITY_ORDER.get(x['severity'].upper(), 3)), 1):
            sev = f['severity'].upper()
            row = t.rows[j]
            row.cells[0].text = sev
            row.cells[1].text = f['rule_id']
            row.cells[2].text = f['name']
            url_short = f['url'][-60:] if len(f['url']) > 60 else f['url']
            row.cells[3].text = url_short
            row.cells[4].text = ', '.join(f.get('compliance', []))

            _set_cell_bg(row.cells[0], _SEV_HEX.get(sev, '64748B'))
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = _rgb(255, 255, 255)
            row.cells[0].paragraphs[0].runs[0].bold = True
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9) if cell.paragraphs[0].runs else Pt(9)

        doc.add_paragraph()

        # Detail per finding
        for j, f in enumerate(sorted(cat_findings, key=lambda x: _SEVERITY_ORDER.get(x['severity'].upper(), 3)), 1):
            _h3(f'Hallazgo {section_num}.{j}: [{f["rule_id"]}] {f["name"]}')

            detail_table = doc.add_table(rows=6, cols=2)
            detail_table.style = 'Table Grid'
            detail_data = [
                ('Severidad', f['severity']),
                ('Regla', f['rule_id']),
                ('URL / Recurso', f['url']),
                ('Línea', str(f.get('line', 'N/A'))),
                ('Compliance', ', '.join(f.get('compliance', []))),
                ('Categoría OWASP', ', '.join(
                    c for c in f.get('compliance', []) if 'OWASP' in c
                ) or 'N/A'),
            ]
            for k, (label, value) in enumerate(detail_data):
                detail_table.rows[k].cells[0].text = label
                detail_table.rows[k].cells[1].text = value
                detail_table.rows[k].cells[0].paragraphs[0].runs[0].bold = True
                _set_cell_bg(detail_table.rows[k].cells[0], 'E8EFF7')
                sev = f['severity'].upper()
                if label == 'Severidad':
                    _set_cell_bg(detail_table.rows[k].cells[1], _SEV_HEX.get(sev, '64748B'))
                    detail_table.rows[k].cells[1].paragraphs[0].runs[0].font.color.rgb = _rgb(255, 255, 255)
                    detail_table.rows[k].cells[1].paragraphs[0].runs[0].bold = True

            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Condición: ').bold = True
            p.add_run(f['description'])

            p = doc.add_paragraph()
            p.add_run('Criterio: ').bold = True
            p.add_run(
                f'Según {", ".join(f.get("compliance", ["OWASP"]))}, este control es obligatorio '
                'para proteger la confidencialidad e integridad de la aplicación.'
            )

            p = doc.add_paragraph()
            p.add_run('Causa probable: ').bold = True
            p.add_run(_probable_cause(f['rule_id']))

            p = doc.add_paragraph()
            p.add_run('Efecto potencial: ').bold = True
            p.add_run(_potential_impact(f['rule_id'], f['severity']))

            doc.add_paragraph()

        section_num += 1

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 5. ANÁLISIS DE BRECHAS OWASP TOP 10
    # ─────────────────────────────────────────────────────────────────────────
    _h1('5. Análisis de Brechas OWASP Top 10 : 2021')
    _para(
        'La siguiente tabla muestra la cobertura de controles respecto a las '
        '10 categorías de riesgo del OWASP Top 10 : 2021.'
    )

    owasp_map = _build_owasp_map(findings)
    owasp_table = doc.add_table(rows=len(_OWASP_CATEGORIES) + 1, cols=4)
    owasp_table.style = 'Table Grid'
    _table_header_row(owasp_table, ['Categoría', 'Descripción', 'Hallazgos', 'Estado'])

    for i, (code, name) in enumerate(_OWASP_CATEGORIES.items(), 1):
        count = owasp_map.get(code, 0)
        status = 'NO CONFORME' if count > 0 else 'CONFORME'
        status_hex = 'DC0000' if count > 0 else '009900'
        row = owasp_table.rows[i]
        row.cells[0].text = f'OWASP-{code}:2021'
        row.cells[1].text = name
        row.cells[2].text = str(count)
        row.cells[3].text = status
        _set_cell_bg(row.cells[3], status_hex)
        row.cells[3].paragraphs[0].runs[0].font.color.rgb = _rgb(255, 255, 255)
        row.cells[3].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 6. MAPA DE RIESGO
    # ─────────────────────────────────────────────────────────────────────────
    _h1('6. Mapa de Riesgo')
    _para(
        'El mapa de riesgo consolida los hallazgos por severidad y categoría, '
        'permitiendo priorizar el esfuerzo de remediación.'
    )

    risk_table = doc.add_table(rows=len(by_category) + 1, cols=6)
    risk_table.style = 'Table Grid'
    _table_header_row(risk_table, ['Categoría', 'Total', 'Crítico', 'Alto', 'Medio', 'Bajo'])

    for i, (cat, cat_f) in enumerate(sorted(by_category.items()), 1):
        c = {'CRITICAL': 0, 'HIGH': 0, 'MEDIUM': 0, 'LOW': 0}
        for f in cat_f:
            sev = f['severity'].upper()
            if sev in c:
                c[sev] += 1
        row = risk_table.rows[i]
        row.cells[0].text = cat
        row.cells[1].text = str(len(cat_f))
        row.cells[2].text = str(c['CRITICAL'])
        row.cells[3].text = str(c['HIGH'])
        row.cells[4].text = str(c['MEDIUM'])
        row.cells[5].text = str(c['LOW'])
        row.cells[0].paragraphs[0].runs[0].bold = True
        if c['CRITICAL']:
            _set_cell_bg(row.cells[2], 'DC0000')
            row.cells[2].paragraphs[0].runs[0].font.color.rgb = _rgb(255, 255, 255)
        if c['HIGH']:
            _set_cell_bg(row.cells[3], 'FF6600')
            row.cells[3].paragraphs[0].runs[0].font.color.rgb = _rgb(255, 255, 255)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 7. CONCLUSIONES
    # ─────────────────────────────────────────────────────────────────────────
    _h1('7. Conclusiones')

    critical_cats = [cat for cat, cat_f in by_category.items()
                     if any(f['severity'] == 'CRITICAL' for f in cat_f)]
    high_cats = [cat for cat, cat_f in by_category.items()
                 if any(f['severity'] == 'HIGH' for f in cat_f)]

    if counts['CRITICAL'] > 0:
        _para(
            f'La auditoría revela un nivel de riesgo CRÍTICO. Se identificaron {counts["CRITICAL"]} hallazgos '
            f'críticos en las categorías: {", ".join(critical_cats)}. Estos hallazgos representan '
            'vulnerabilidades actualmente explotables que podrían comprometer la confidencialidad, '
            'integridad y disponibilidad de la aplicación y sus usuarios.',
            bold=False
        )

    if counts['HIGH'] > 0:
        _para(
            f'Adicionalmente, {counts["HIGH"]} hallazgos de severidad ALTA en las categorías '
            f'{", ".join(high_cats)} requieren remediación dentro de 7 días.'
        )

    _para(
        f'Se recomienda iniciar un ciclo de remediación inmediato para los hallazgos CRÍTICOS y ALTOS, '
        f'seguido de una nueva auditoría de verificación en un plazo no mayor a 30 días.'
    )

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 8. RECOMENDACIONES PRIORIZADAS
    # ─────────────────────────────────────────────────────────────────────────
    _h1('8. Recomendaciones Priorizadas')

    rec_num = 1
    for sev_filter in ('CRITICAL', 'HIGH', 'MEDIUM', 'LOW'):
        filtered = [f for f in findings if f['severity'].upper() == sev_filter]
        if not filtered:
            continue
        _h2(f'Severidad {sev_filter} ({len(filtered)} hallazgos)')
        cats_in_sev = list(dict.fromkeys(_get_category(f['rule_id']) for f in filtered))
        for cat in cats_in_sev:
            _h3(f'{rec_num}. Remediación: {cat}')
            _para(_REMEDIATION_GUIDES.get(cat, 'Revisar la documentación del estándar de referencia.'))
            cat_findings = [f for f in filtered if _get_category(f['rule_id']) == cat]
            _para(f'Afecta a {len(cat_findings)} hallazgo(s):', bold=True)
            for f in cat_findings[:5]:
                url_short = f['url'][-70:] if len(f['url']) > 70 else f['url']
                _para(f'  • [{f["rule_id"]}] {url_short}')
            if len(cat_findings) > 5:
                _para(f'  ... y {len(cat_findings) - 5} más.')
            rec_num += 1

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 9. PLAN DE REMEDIACIÓN CON SLAs
    # ─────────────────────────────────────────────────────────────────────────
    _h1('9. Plan de Remediación con SLAs')
    _para(
        'La siguiente tabla define los responsables, plazos y criterios de verificación '
        'para cada hallazgo activo.'
    )

    # Group by category, show top 20
    shown = findings[:20]
    plan_table = doc.add_table(rows=len(shown) + 1, cols=6)
    plan_table.style = 'Table Grid'
    _table_header_row(plan_table, ['Regla', 'Severidad', 'Categoría', 'SLA', 'Responsable', 'Estado'])

    for i, f in enumerate(shown, 1):
        sev = f['severity'].upper()
        row = plan_table.rows[i]
        row.cells[0].text = f['rule_id']
        row.cells[1].text = sev
        row.cells[2].text = _get_category(f['rule_id'])
        row.cells[3].text = _SLA_BY_SEVERITY.get(sev, '30 días')
        row.cells[4].text = '[Responsable técnico]'
        row.cells[5].text = 'PENDIENTE'
        _set_cell_bg(row.cells[1], _SEV_HEX.get(sev, '64748B'))
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = _rgb(255, 255, 255)
        row.cells[1].paragraphs[0].runs[0].bold = True
        _set_cell_bg(row.cells[5], 'CC9900')
        row.cells[5].paragraphs[0].runs[0].font.color.rgb = _rgb(255, 255, 255)

    if len(findings) > 20:
        _para(f'\n... y {len(findings) - 20} hallazgos adicionales no mostrados en esta tabla.')

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 10. ANEXOS
    # ─────────────────────────────────────────────────────────────────────────
    _h1('10. Anexos')

    _h2('10.1 Payloads de Prueba Utilizados')
    _para('Los siguientes payloads no destructivos fueron utilizados para detección:')
    payloads = [
        ('XSS', "<script>alert(1)</script>, \"><script>alert(1)</script>"),
        ('SQLi', "' OR '1'='1"),
        ('Open Redirect', 'https://evil.com'),
        ('CORS', 'Origin: https://evil.com'),
    ]
    for ptype, payload in payloads:
        p = doc.add_paragraph()
        p.add_run(f'{ptype}: ').bold = True
        p.add_run(payload)

    _h2('10.2 Glosario')
    glossary = [
        ('DAST', 'Dynamic Application Security Testing — pruebas sobre la aplicación en ejecución.'),
        ('SAST', 'Static Application Security Testing — análisis del código fuente sin ejecución.'),
        ('CSP', 'Content Security Policy — mecanismo de control de recursos cargables.'),
        ('HSTS', 'HTTP Strict Transport Security — fuerza conexiones HTTPS.'),
        ('CSRF', 'Cross-Site Request Forgery — ataque de falsificación de petición entre sitios.'),
        ('XSS', 'Cross-Site Scripting — inyección de scripts maliciosos en el navegador víctima.'),
        ('SQLi', 'SQL Injection — manipulación de consultas SQL a través de input de usuario.'),
        ('CORS', 'Cross-Origin Resource Sharing — política de control de acceso entre dominios.'),
        ('OWASP', 'Open Worldwide Application Security Project — organización de referencia en seguridad web.'),
    ]
    gls_table = doc.add_table(rows=len(glossary) + 1, cols=2)
    gls_table.style = 'Table Grid'
    _table_header_row(gls_table, ['Término', 'Definición'])
    for i, (term, defn) in enumerate(glossary, 1):
        gls_table.rows[i].cells[0].text = term
        gls_table.rows[i].cells[1].text = defn
        gls_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    _h2('10.3 Declaración de Autorización')
    _para(
        f'Esta auditoría fue realizada con autorización explícita de {authorized_by} de la empresa {empresa}. '
        'El acceso a los sistemas fue estrictamente pasivo y/o con payloads no destructivos. '
        f'Fecha de autorización: {fecha}. '
        'AuditLens no se hace responsable del uso no autorizado de este informe o de las técnicas descritas.'
    )

    # Save
    doc.save(output_path)
    abs_path = os.path.abspath(output_path)
    print(f'\033[92m[AuditLens]\033[0m Informe Word guardado: \033[1m{abs_path}\033[0m')
    return abs_path


# ── Helpers for finding descriptions ─────────────────────────────────────────

def _probable_cause(rule_id: str) -> str:
    causes = {
        'WEB-SSL': 'Configuración del servidor web sin política de ciphers restrictiva o certificado no renovado.',
        'WEB-HDR': 'Cabeceras de seguridad no configuradas en el servidor web o el framework de la aplicación.',
        'WEB-CKI': 'Framework o lenguaje sin configuración explícita de atributos de cookie en producción.',
        'WEB-CRS': 'Política CORS configurada de forma permisiva para facilitar el desarrollo, no ajustada para producción.',
        'WEB-FRM': 'Formulario generado sin integración del mecanismo CSRF del framework o sin validación de acción HTTPS.',
        'WEB-XSS': 'Input de usuario renderizado en HTML sin codificación previa (output encoding).',
        'WEB-SQL': 'Consultas SQL construidas por concatenación de strings sin uso de prepared statements.',
        'WEB-RDR': 'Parámetro de redirección sin validación de dominio destino.',
        'WEB-PTH': 'Archivo sensible accesible en DocumentRoot sin restricción de acceso en el servidor web.',
        'WEB-INF': 'Modo DEBUG activo en producción o páginas de error genéricas no configuradas.',
    }
    for prefix, cause in causes.items():
        if rule_id.startswith(prefix):
            return cause
    return 'Ausencia de control de seguridad requerido por el estándar.'


def _potential_impact(rule_id: str, severity: str) -> str:
    impacts = {
        'WEB-SSL': 'Intercepción del tráfico cifrado (MITM), downgrade de TLS, exposición de credenciales en tránsito.',
        'WEB-HDR': 'XSS, clickjacking, MIME sniffing y filtración de datos via Referer.',
        'WEB-CKI': 'Robo de sesión via XSS o red no cifrada; ataques CSRF.',
        'WEB-CRS': 'Robo de datos de la API por sitios maliciosos en el navegador del usuario.',
        'WEB-FRM': 'Ejecución de acciones no autorizadas en nombre del usuario autenticado (CSRF).',
        'WEB-XSS': 'Robo de cookies de sesión, redirección a sitios de phishing, defacement.',
        'WEB-SQL': 'Exfiltración de la base de datos, bypass de autenticación, destrucción de datos.',
        'WEB-RDR': 'Redirección de usuarios a sitios de phishing en ataques de ingeniería social.',
        'WEB-PTH': 'Exposición de credenciales (.env), código fuente (.git), backups de base de datos.',
        'WEB-INF': 'Fingerprinting de la tecnología, facilitación de ataques dirigidos a versiones específicas.',
    }
    for prefix, impact in impacts.items():
        if rule_id.startswith(prefix):
            return impact
    return 'Exposición de información sensible o habilitación de vectores de ataque adicionales.'


def _build_owasp_map(findings: List[dict]) -> Dict[str, int]:
    owasp_counts: Dict[str, int] = {}
    for f in findings:
        for c in f.get('compliance', []):
            m = re.search(r'OWASP-A(\d{2}):', c)
            if m:
                code = f'A{m.group(1)}'
                owasp_counts[code] = owasp_counts.get(code, 0) + 1
    return owasp_counts


import re
