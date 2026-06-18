"""
AuditLens — Generador de Reporte Ley 21.719 (Chile)

Genera reporte HTML standalone y documento Word listo para
presentar a la Agencia de Protección de Datos Personales (APDP).
"""

from __future__ import annotations

import json
import os
from datetime import date
from typing import Any, Dict, List

from .ley21719_mapper import LEY_ARTICULOS, calculate_compliance_score, enrich_with_ley21719


_SEV_COLOR = {
    'CRITICAL': '#da3633', 'HIGH': '#e3b341',
    'MEDIUM':   '#388bfd', 'LOW':  '#3fb950',
}

_RISK_COLOR = {
    'CRÍTICO': '#da3633', 'ALTO': '#e3b341',
    'MEDIO':   '#388bfd', 'BAJO': '#3fb950',
}


def generate_ley21719_html(
    findings: List[dict],
    score_data: Dict[str, Any],
    output_path: str,
    empresa: str = 'Empresa',
    rut_empresa: str = '',
    auditor: str = '',
) -> str:
    """Generate standalone HTML compliance report for Ley 21.719."""

    score      = score_data['score']
    risk       = score_data['risk_level']
    risk_color = _RISK_COLOR.get(risk, '#8b949e')
    score_color = '#da3633' if score < 40 else '#e3b341' if score < 70 else '#3fb950'
    today      = date.today().strftime('%d/%m/%Y')

    # Score gauge SVG
    circ_pct   = score / 100
    dash_val   = round(circ_pct * 283)
    gauge_svg  = f"""
    <svg viewBox="0 0 100 100" class="score-gauge">
      <circle cx="50" cy="50" r="45" fill="none" stroke="#21262d" stroke-width="10"/>
      <circle cx="50" cy="50" r="45" fill="none" stroke="{score_color}" stroke-width="10"
        stroke-dasharray="{dash_val} 283" stroke-dashoffset="71"
        stroke-linecap="round" transform="rotate(-90 50 50)"/>
      <text x="50" y="46" text-anchor="middle" fill="{score_color}"
        font-size="22" font-weight="bold">{score}</text>
      <text x="50" y="62" text-anchor="middle" fill="#8b949e" font-size="10">/100</text>
    </svg>"""

    # Articles violated table
    arts_rows = ''
    for art, detail in score_data.get('breakdown', {}).items():
        info = LEY_ARTICULOS.get(art, {})
        arts_rows += f"""
        <tr>
          <td class="mono" style="color:#f78166">{art}</td>
          <td>{info.get('titulo','')}</td>
          <td style="font-size:11px;color:#8b949e">{info.get('obligacion','')}</td>
          <td style="color:#da3633;text-align:right">-{detail.get('deduction',0)} pts</td>
        </tr>"""

    # PII findings table
    pii_findings = [f for f in findings if f.get('rule_id','').startswith('PII-')]
    pii_rows = ''
    for f in sorted(pii_findings, key=lambda x: {'CRITICAL':0,'HIGH':1,'MEDIUM':2,'LOW':3}.get(x.get('severity','LOW'),4)):
        sev   = f.get('severity','LOW')
        color = _SEV_COLOR.get(sev,'#8b949e')
        arts  = ', '.join(f.get('ley21719',[]))
        pii_rows += f"""
        <tr>
          <td><span class="sev-dot" style="background:{color}"></span>{sev}</td>
          <td>{f.get('name','')}</td>
          <td class="mono">{f.get('file','').split('/')[-1]}:{f.get('line','')}</td>
          <td style="font-size:11px;color:#79c0ff">{arts}</td>
          <td class="mono" style="color:#8b949e;font-size:11px">{f.get('pii_value_preview','')}</td>
        </tr>"""

    # Security findings table
    sec_findings = [f for f in findings if not f.get('rule_id','').startswith('PII-')]
    sec_rows = ''
    for f in sorted(sec_findings, key=lambda x: {'CRITICAL':0,'HIGH':1,'MEDIUM':2,'LOW':3}.get(x.get('severity','LOW'),4))[:30]:
        sev   = f.get('severity','LOW')
        color = _SEV_COLOR.get(sev,'#8b949e')
        arts  = ', '.join(f.get('ley21719',[]))
        sec_rows += f"""
        <tr>
          <td><span class="sev-dot" style="background:{color}"></span>{sev}</td>
          <td>{f.get('rule_id','')}</td>
          <td>{f.get('name','')[:50]}</td>
          <td class="mono">{f.get('file','').split('/')[-1]}:{f.get('line','')}</td>
          <td style="font-size:11px;color:#79c0ff">{arts}</td>
        </tr>"""

    # EIPD checklist
    eipd_items = [
        ('Identificar datos personales que se tratan', bool(pii_findings)),
        ('Documentar finalidad del tratamiento', False),
        ('Evaluar base legal para cada tratamiento', False),
        ('Implementar medidas técnicas de seguridad', score >= 70),
        ('Verificar mecanismos ARCOP (acceso, rectificación, etc.)', False),
        ('Designar Encargado de Protección de Datos (si aplica)', False),
        ('Establecer protocolo de notificación de brechas (72h)', False),
        ('Revisar contratos con terceros que traten datos', False),
        ('Capacitar al personal en protección de datos', False),
        ('Mantener registro de actividades de tratamiento', False),
    ]
    eipd_rows = ''
    for item, done in eipd_items:
        icon  = '✓' if done else '○'
        color = '#3fb950' if done else '#6e7681'
        eipd_rows += f'<tr><td style="color:{color};font-size:16px">{icon}</td><td>{item}</td><td style="color:{color}">' + ('Completado' if done else 'Pendiente') + '</td></tr>'

    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>Informe Cumplimiento Ley 21.719 — {empresa}</title>
<style>
* {{ box-sizing:border-box; margin:0; padding:0; }}
body {{ font-family:'Segoe UI',system-ui,sans-serif; background:#0d1117; color:#e6edf3; padding:32px; max-width:1100px; margin:0 auto; }}
h1 {{ color:#58a6ff; font-size:22px; margin-bottom:4px; }}
h2 {{ color:#58a6ff; font-size:13px; text-transform:uppercase; letter-spacing:1px; margin:0; }}
.subtitle {{ color:#8b949e; font-size:13px; margin-bottom:28px; }}
.header-grid {{ display:grid; grid-template-columns:1fr auto; gap:24px; align-items:start; margin-bottom:28px; }}
.meta-grid {{ display:grid; grid-template-columns:1fr 1fr; gap:8px; font-size:13px; }}
.meta-row {{ display:flex; gap:8px; }}
.meta-label {{ color:#8b949e; min-width:120px; }}

.score-gauge {{ width:160px; height:160px; }}
.score-box {{ text-align:center; }}
.risk-badge {{ display:inline-block; padding:4px 14px; border-radius:20px; font-size:14px; font-weight:700;
  background:{risk_color}22; border:1px solid {risk_color}; color:{risk_color}; margin-top:8px; }}

.stats-grid {{ display:grid; grid-template-columns:repeat(4,1fr); gap:12px; margin-bottom:24px; }}
.stat-card {{ background:#161b22; border:1px solid #30363d; border-radius:8px; padding:16px; }}
.stat-value {{ font-size:28px; font-weight:700; }}
.stat-label {{ font-size:11px; color:#8b949e; text-transform:uppercase; letter-spacing:1px; margin-top:4px; }}

.section {{ background:#161b22; border:1px solid #30363d; border-radius:8px; padding:20px; margin-bottom:16px; }}
.section-header {{ display:flex; justify-content:space-between; align-items:center; margin-bottom:14px; }}
.multa-banner {{ background:#3d1a1a; border:1px solid #da3633; border-radius:8px; padding:14px 20px; margin-bottom:16px; color:#f85149; font-size:13px; }}

table {{ width:100%; border-collapse:collapse; font-size:12px; }}
th {{ text-align:left; padding:8px; color:#8b949e; border-bottom:1px solid #30363d; font-size:11px; text-transform:uppercase; }}
td {{ padding:7px 8px; border-bottom:1px solid #21262d; vertical-align:top; }}
.mono {{ font-family:monospace; font-size:11px; }}
.sev-dot {{ width:8px; height:8px; border-radius:50%; display:inline-block; margin-right:6px; }}

.watermark {{ text-align:center; color:#30363d; font-size:11px; margin-top:32px; }}
</style>
</head>
<body>

<div class="header-grid">
  <div>
    <h1>🔒 Informe de Cumplimiento — Ley 21.719</h1>
    <p class="subtitle">Protección de Datos Personales | República de Chile</p>
    <div class="meta-grid">
      <div class="meta-row"><span class="meta-label">Empresa:</span> <strong>{empresa}</strong></div>
      <div class="meta-row"><span class="meta-label">RUT Empresa:</span> {rut_empresa or '—'}</div>
      <div class="meta-row"><span class="meta-label">Auditor:</span> {auditor or '—'}</div>
      <div class="meta-row"><span class="meta-label">Fecha:</span> {today}</div>
      <div class="meta-row"><span class="meta-label">Hallazgos totales:</span> {len(findings)}</div>
      <div class="meta-row"><span class="meta-label">PII detectado:</span> {len(pii_findings)}</div>
    </div>
  </div>
  <div class="score-box">
    {gauge_svg}
    <div class="risk-badge">{risk}</div>
    <div style="font-size:11px;color:#8b949e;margin-top:6px">Nivel de riesgo</div>
  </div>
</div>

<div class="stats-grid">
  <div class="stat-card">
    <div class="stat-value" style="color:#da3633">{score_data['pii_findings']}</div>
    <div class="stat-label">Datos personales expuestos</div>
  </div>
  <div class="stat-card">
    <div class="stat-value" style="color:#e3b341">{len(score_data['articles_violated'])}</div>
    <div class="stat-label">Artículos incumplidos</div>
  </div>
  <div class="stat-card">
    <div class="stat-value" style="color:#388bfd">{score_data['security_findings']}</div>
    <div class="stat-label">Vulnerabilidades técnicas</div>
  </div>
  <div class="stat-card">
    <div class="stat-value" style="color:{score_color}">{score}</div>
    <div class="stat-label">Score de cumplimiento</div>
  </div>
</div>

<div class="multa-banner">
  ⚠️ <strong>Estimación de riesgo de multa:</strong> {score_data['multa_estimada']}
  <br><span style="font-size:11px;color:#8b949e">Basado en Art. 49 Ley 21.719. Esta estimación no constituye asesoramiento legal.</span>
</div>

<div class="section">
  <div class="section-header"><h2>Artículos de la Ley 21.719 incumplidos</h2></div>
  <table>
    <thead><tr><th>Artículo</th><th>Materia</th><th>Obligación</th><th>Impacto</th></tr></thead>
    <tbody>{arts_rows if arts_rows else '<tr><td colspan="4" style="color:#6e7681;text-align:center;padding:16px">No se detectaron artículos incumplidos.</td></tr>'}</tbody>
  </table>
</div>

<div class="section">
  <div class="section-header">
    <h2>Datos Personales detectados en código</h2>
    <span style="color:#8b949e;font-size:12px">{len(pii_findings)} hallazgos</span>
  </div>
  <table>
    <thead><tr><th>Severidad</th><th>Tipo de dato</th><th>Archivo</th><th>Artículos</th><th>Valor (enmascarado)</th></tr></thead>
    <tbody>{pii_rows if pii_rows else '<tr><td colspan="5" style="color:#6e7681;text-align:center;padding:16px">No se detectaron datos personales expuestos.</td></tr>'}</tbody>
  </table>
</div>

<div class="section">
  <div class="section-header">
    <h2>Vulnerabilidades de seguridad técnica</h2>
    <span style="color:#8b949e;font-size:12px">{len(sec_findings)} hallazgos (top 30)</span>
  </div>
  <table>
    <thead><tr><th>Severidad</th><th>Regla</th><th>Descripción</th><th>Archivo</th><th>Artículos</th></tr></thead>
    <tbody>{sec_rows if sec_rows else '<tr><td colspan="5" style="color:#6e7681;text-align:center;padding:16px">No se detectaron vulnerabilidades de seguridad.</td></tr>'}</tbody>
  </table>
</div>

<div class="section">
  <div class="section-header">
    <h2>Lista de verificación EIPD (Evaluación de Impacto)</h2>
    <span style="color:#8b949e;font-size:12px">Art. 28 Ley 21.719</span>
  </div>
  <table>
    <thead><tr><th style="width:40px"></th><th>Ítem</th><th>Estado</th></tr></thead>
    <tbody>{eipd_rows}</tbody>
  </table>
</div>

<div class="watermark">
  Generado por AuditLens · Informe de cumplimiento Ley 21.719 · {today}
  <br>Este informe es de carácter técnico. Para efectos legales, consulte a un profesional habilitado.
</div>
</body>
</html>"""

    with open(output_path, 'w', encoding='utf-8') as fh:
        fh.write(html)
    print(f'\033[92m[AuditLens Ley 21.719]\033[0m Reporte HTML generado: {output_path}')
    return output_path


def generate_ley21719_docx(
    findings: List[dict],
    score_data: Dict[str, Any],
    output_path: str,
    empresa: str = 'Empresa',
    rut_empresa: str = '',
    auditor: str = '',
    rut_auditor: str = '',
) -> str:
    """Generate Word document for Ley 21.719 compliance report."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print('\033[91m[AuditLens]\033[0m python-docx no instalado: pip install python-docx')
        return ''

    today    = date.today().strftime('%d de %B de %Y').replace(
        'January','enero').replace('February','febrero').replace('March','marzo')
    score    = score_data['score']
    risk     = score_data['risk_level']
    pii_f    = [f for f in findings if f.get('rule_id','').startswith('PII-')]
    sec_f    = [f for f in findings if not f.get('rule_id','').startswith('PII-')]

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    def _heading(text, level=1):
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = RGBColor(0x1f, 0x6f, 0xeb)
        return p

    def _add_table(headers, rows, col_widths=None):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        for row_data in rows:
            row = table.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = str(val)
                row[i].paragraphs[0].runs[0].font.size = Pt(9)
        return table

    # ── Portada ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run('INFORME DE CUMPLIMIENTO\nLEY N° 21.719\nPROTECCIÓN DE DATOS PERSONALES')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1f, 0x6f, 0xeb)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Empresa: {empresa}\n').bold = True
    if rut_empresa:
        meta.add_run(f'RUT: {rut_empresa}\n')
    meta.add_run(f'Fecha: {today}\n')
    if auditor:
        meta.add_run(f'Auditor: {auditor}')
        if rut_auditor:
            meta.add_run(f' · RUT: {rut_auditor}')

    doc.add_paragraph()
    score_p = doc.add_paragraph()
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = score_p.add_run(f'Score de Cumplimiento: {score}/100 — Riesgo {risk}')
    sr.bold = True
    sr.font.size = Pt(14)
    color_map = {'CRÍTICO': RGBColor(0xda,0x36,0x33), 'ALTO': RGBColor(0xe3,0xb3,0x41),
                 'MEDIO': RGBColor(0x38,0x8b,0xfd), 'BAJO': RGBColor(0x3f,0xb9,0x50)}
    sr.font.color.rgb = color_map.get(risk, RGBColor(0x8b,0x94,0x9e))

    doc.add_page_break()

    # ── 1. Resumen ejecutivo ──────────────────────────────────────────────────
    _heading('1. Resumen Ejecutivo', 1)
    doc.add_paragraph(
        f'El presente informe corresponde a la evaluación de cumplimiento de la Ley N° 21.719 '
        f'sobre Protección de Datos Personales de la empresa {empresa}, realizada el {today}. '
        f'Se analizó el código fuente y la configuración de los sistemas de información '
        f'en busca de datos personales expuestos y vulnerabilidades de seguridad técnica.'
    )
    doc.add_paragraph()
    doc.add_paragraph(f'Score de cumplimiento obtenido: {score}/100 (Riesgo {risk})')
    doc.add_paragraph(f'Total hallazgos: {len(findings)} ({len(pii_f)} datos personales, {len(sec_f)} vulnerabilidades técnicas)')
    doc.add_paragraph(f'Artículos incumplidos: {", ".join(score_data.get("articles_violated", []))}')
    doc.add_paragraph(f'Estimación de multa: {score_data["multa_estimada"]}')

    # ── 2. Marco legal ────────────────────────────────────────────────────────
    _heading('2. Marco Legal', 1)
    doc.add_paragraph(
        'La Ley N° 21.719, publicada el 13 de diciembre de 2023 en el Diario Oficial, '
        'establece un nuevo marco para la protección de datos personales en Chile, '
        'reemplazando la Ley N° 19.628 de 1999. La ley otorgó un plazo de 24 meses '
        'para adecuarse, con vigencia plena a partir de diciembre de 2025.'
    )

    # ── 3. Artículos incumplidos ──────────────────────────────────────────────
    _heading('3. Artículos Incumplidos', 1)
    art_rows = []
    for art in score_data.get('articles_violated', []):
        info = LEY_ARTICULOS.get(art, {})
        ded  = score_data['breakdown'].get(art, {}).get('deduction', 0)
        art_rows.append([art, info.get('titulo',''), info.get('obligacion',''), f'-{ded} pts'])
    if art_rows:
        _add_table(['Artículo', 'Materia', 'Obligación', 'Impacto'], art_rows)
    else:
        doc.add_paragraph('No se detectaron artículos incumplidos.')

    # ── 4. Datos personales detectados ────────────────────────────────────────
    _heading('4. Datos Personales Detectados en Código Fuente', 1)
    doc.add_paragraph(
        f'Se detectaron {len(pii_f)} instancias de datos personales expuestos en el '
        f'código fuente. Esto constituye una infracción directa al principio de seguridad '
        f'(Art. 14) y las medidas técnicas obligatorias (Art. 14d) de la Ley 21.719.'
    )
    if pii_f:
        pii_rows_docx = [
            [
                f.get('severity',''),
                f.get('name',''),
                f'{f.get("file","").split("/")[-1]}:{f.get("line","")}',
                ', '.join(f.get('ley21719',[])),
                f.get('pii_value_preview',''),
            ]
            for f in sorted(pii_f, key=lambda x: {'CRITICAL':0,'HIGH':1,'MEDIUM':2,'LOW':3}.get(x.get('severity','LOW'),4))
        ]
        _add_table(['Severidad','Tipo de Dato','Archivo:Línea','Artículos','Valor (enmascarado)'], pii_rows_docx)

    # ── 5. Vulnerabilidades técnicas ──────────────────────────────────────────
    _heading('5. Vulnerabilidades de Seguridad Técnica', 1)
    doc.add_paragraph(
        f'Se detectaron {len(sec_f)} vulnerabilidades de seguridad técnica que podrían '
        f'comprometer la integridad y confidencialidad de los datos personales tratados.'
    )
    if sec_f:
        sec_rows_docx = [
            [
                f.get('severity',''),
                f.get('rule_id',''),
                f.get('name','')[:60],
                f'{f.get("file","").split("/")[-1]}:{f.get("line","")}',
                ', '.join(f.get('ley21719',[])),
            ]
            for f in sorted(sec_f, key=lambda x: {'CRITICAL':0,'HIGH':1,'MEDIUM':2,'LOW':3}.get(x.get('severity','LOW'),4))[:25]
        ]
        _add_table(['Severidad','Regla','Descripción','Archivo','Artículos'], sec_rows_docx)

    # ── 6. EIPD ──────────────────────────────────────────────────────────────
    _heading('6. Evaluación de Impacto en la Protección de Datos (EIPD)', 1)
    doc.add_paragraph(
        'Conforme al Art. 28 de la Ley 21.719, se presenta la siguiente lista de '
        'verificación para la Evaluación de Impacto:'
    )
    eipd_items = [
        'Identificar y documentar todos los datos personales tratados',
        'Documentar la finalidad y base legal de cada tratamiento',
        'Evaluar la proporcionalidad del tratamiento',
        'Implementar medidas técnicas de seguridad (cifrado, control de acceso)',
        'Implementar mecanismos para el ejercicio de derechos ARCOP',
        'Designar Encargado de Protección de Datos (si aplica)',
        'Establecer protocolo de notificación de brechas en 72 horas',
        'Revisar y actualizar contratos con terceros encargados del tratamiento',
        'Capacitar al personal en materia de protección de datos',
        'Mantener registro actualizado de actividades de tratamiento',
    ]
    for item in eipd_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'☐ {item}')

    # ── 7. Recomendaciones ────────────────────────────────────────────────────
    _heading('7. Recomendaciones', 1)
    recs = [
        ('INMEDIATO', 'Eliminar todos los datos personales hardcodeados del código fuente y repositorios git.'),
        ('INMEDIATO', 'Rotar credenciales y secretos expuestos en historial de commits.'),
        ('CORTO PLAZO', 'Implementar cifrado en tránsito (TLS) y en reposo para bases de datos con datos personales.'),
        ('CORTO PLAZO', 'Eliminar modo DEBUG en ambientes de producción.'),
        ('MEDIANO PLAZO', 'Implementar gestión de consentimiento con registro auditable.'),
        ('MEDIANO PLAZO', 'Desarrollar procedimientos para el ejercicio de derechos ARCOP.'),
        ('MEDIANO PLAZO', 'Realizar la Evaluación de Impacto (EIPD) formalmente.'),
        ('LARGO PLAZO', 'Designar Encargado de Protección de Datos (DPO) si aplica.'),
        ('LARGO PLAZO', 'Implementar programa de capacitación continua en protección de datos.'),
    ]
    for plazo, rec in recs:
        p = doc.add_paragraph()
        r1 = p.add_run(f'[{plazo}] ')
        r1.bold = True
        p.add_run(rec)

    # ── 8. Declaración ────────────────────────────────────────────────────────
    _heading('8. Declaración del Auditor', 1)
    doc.add_paragraph(
        f'El presente informe ha sido elaborado mediante análisis automatizado de código fuente '
        f'y configuraciones de sistemas. Los hallazgos reflejan el estado técnico de los sistemas '
        f'analizados a la fecha {today}. Este informe es de carácter técnico y no reemplaza '
        f'la asesoría legal especializada.'
    )
    doc.add_paragraph()
    doc.add_paragraph(f'Auditor: {auditor or "___________________________"}')
    if rut_auditor:
        doc.add_paragraph(f'RUT: {rut_auditor}')
    doc.add_paragraph(f'Fecha: {today}')
    doc.add_paragraph('Firma: ___________________________')

    doc.save(output_path)
    print(f'\033[92m[AuditLens Ley 21.719]\033[0m Documento Word generado: {output_path}')
    return output_path
